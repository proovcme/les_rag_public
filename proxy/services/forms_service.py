"""W11.3/W19 — сервис типовых форм документов: дескриптор + данные → документ.

Канон ADR-11 («документы = шаблоны + данные», статья о границах LLM): форма хранится
как машиночитаемый **дескриптор** (поля/типы/источники/правовая основа), значения
**резолвятся детерминированно из данных объекта** (W17/W8/W11/W17.2), а документ —
всегда ВЫХОД, генерируется рендерером в нужном формате. LLM при заполнении не участвует
и не трогает числа/XML. Шаблон-вёрстка (фирменный бланк/штамп) хранится в родном формате
образца; если файла нет — рендер строит документ из дескриптора (template-less fallback).

Слои:
  1. registry   — загрузка дескрипторов из config/forms/*.yaml
  2. resolve    — значения полей из данных объекта (0 LLM): project.* / field.* / edges.* / manual / const / date
  3. render[fmt]— docx (фокус волны) / xlsx / html — за единым интерфейсом
"""
from __future__ import annotations

import time
from pathlib import Path
from typing import Any

FORMS_DIR = Path("config/forms")
OUTPUT_DIR = Path("data/forms_out")

SUPPORTED_FORMATS = ("docx", "xlsx", "html")


# ── 1. реестр дескрипторов ───────────────────────────────────────────

def _forms_dir() -> Path:
    import os
    return Path(os.getenv("LES_FORMS_DIR", str(FORMS_DIR)))


def load_descriptor(form_id: str) -> dict[str, Any] | None:
    import yaml
    path = _forms_dir() / f"{form_id}.yaml"
    if not path.is_file():
        return None
    with path.open(encoding="utf-8") as fh:
        data = yaml.safe_load(fh) or {}
    data.setdefault("id", form_id)
    data.setdefault("fields", [])
    return data


def list_forms() -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for p in sorted(_forms_dir().glob("*.yaml")):
        d = load_descriptor(p.stem) or {}
        out.append({
            "id": d.get("id"),
            "title": d.get("title", d.get("id")),
            "legal_basis": d.get("legal_basis", ""),
            "fields": len(d.get("fields", [])),
            "formats": [f for f in SUPPORTED_FORMATS],
        })
    return out


# ── 2. детерминированный резолв полей (0 LLM) ─────────────────────────

def _fmt_value(value: Any, ftype: str) -> str:
    if value is None:
        return ""
    if ftype == "list" and isinstance(value, (list, tuple)):
        return ", ".join(str(v) for v in value)
    if ftype == "num":
        try:
            num = float(value)
            return f"{num:g}"
        except (TypeError, ValueError):
            return str(value)
    return str(value)


def _resolve_source(source: str, project_id: int | None, manual: dict[str, Any], field_key: str) -> Any:
    """Один источник → значение. Никакого LLM: только SQL/сервисы/справочники."""
    source = (source or "").strip()
    if source == "manual":
        return manual.get(field_key, "")
    if source == "date.today":
        return time.strftime("%d.%m.%Y", time.localtime())
    if source.startswith("const:"):
        return source[len("const:"):]

    head, _, tail = source.partition(".")
    if head == "project" and project_id:
        from proxy.services.project_service import get_project
        proj = get_project(project_id) or {}
        return proj.get(tail, "")
    if head == "field":
        # field.total  или  field.total@<захватка>
        from proxy.services.field_intake_service import aggregate_volumes
        metric, _, zah = tail.partition("@")
        try:
            rows = aggregate_volumes(status="confirmed", zahvatka=zah) if zah else aggregate_volumes(status="confirmed")
        except TypeError:
            rows = aggregate_volumes(status="confirmed")
        except Exception:
            rows = []
        if metric == "total":
            return round(sum(float(r.get("total") or 0) for r in rows), 3)
        return ""
    if head == "edges":
        # edges.references_ntd → уникальные цели рёбер этого типа
        from proxy.services.edge_service import list_edges
        edge_type = tail
        seen: list[str] = []
        for e in list_edges(5000):
            if e.get("edge_type") == edge_type and e.get("dst_id") not in seen:
                seen.append(e.get("dst_id"))
        return seen
    return ""


def resolve_fields(form_id: str, project_id: int | None = None, manual: dict[str, Any] | None = None) -> dict[str, Any] | None:
    """Все поля формы с разрешёнными значениями. `needs_input=True` — ручной ввод не задан."""
    descriptor = load_descriptor(form_id)
    if descriptor is None:
        return None
    manual = manual or {}
    fields = []
    for f in descriptor.get("fields", []):
        ftype = f.get("type", "str")
        source = f.get("source", "manual")
        raw = _resolve_source(source, project_id, manual, f["key"])
        value = _fmt_value(raw, ftype)
        fields.append({
            "key": f["key"],
            "label": f.get("label", f["key"]),
            "type": ftype,
            "source": source,
            "value": value,
            "needs_input": source == "manual" and not value,
        })
    columns = list(descriptor.get("columns", []) or [])
    return {
        "id": descriptor["id"],
        "title": descriptor.get("title", descriptor["id"]),
        "legal_basis": descriptor.get("legal_basis", ""),
        "fields": fields,
        "columns": columns,
        "rows": _resolve_table_rows(descriptor, columns) if columns else [],
    }


def _resolve_table_rows(descriptor: dict[str, Any], columns: list[str]) -> list[list[str]]:
    """Строки табличной формы. mode=blank → N пустых строк (бланк по ГОСТ).

    Заполнение из данных объекта (spec/bor/smeta) — отдельный шаг; сейчас отдаём бланк
    с канонической шапкой колонок (ГОСТ), чтобы форма была собрана и печатаема.
    """
    table = descriptor.get("table", {}) or {}
    count = int(table.get("rows", 12) or 12)
    return [["" for _ in columns] for _ in range(max(0, count))]


# ── 3. рендереры (за единым интерфейсом) ─────────────────────────────

def _output_dir() -> Path:
    import os
    d = Path(os.getenv("LES_FORMS_OUT_DIR", str(OUTPUT_DIR)))
    d.mkdir(parents=True, exist_ok=True)
    return d


def _value_map(resolved: dict[str, Any]) -> dict[str, str]:
    return {f["key"]: f["value"] for f in resolved["fields"]}


def render_html(resolved: dict[str, Any]) -> str:
    from html import escape
    rows = "".join(
        f"<tr><td style='padding:4px 8px;color:#555'>{escape(f['label'])}</td>"
        f"<td style='padding:4px 8px'><b>{escape(f['value']) or '—'}</b></td></tr>"
        for f in resolved["fields"]
    )
    table_html = ""
    cols = resolved.get("columns") or []
    if cols:
        head = "".join(f"<th style='border:1px solid #999;padding:3px 6px'>{escape(c)}</th>" for c in cols)
        body = "".join(
            "<tr>" + "".join(
                f"<td style='border:1px solid #ccc;padding:3px 6px'>{escape(c)}&nbsp;</td>" for c in r
            ) + "</tr>"
            for r in resolved.get("rows", [])
        )
        table_html = (
            f"<table style='border-collapse:collapse;margin-top:10px;font-size:.8em'>"
            f"<tr>{head}</tr>{body}</table>"
        )
    return (
        f"<h3>{escape(resolved['title'])}</h3>"
        f"<div style='font-size:.8em;color:#777'>{escape(resolved['legal_basis'])}</div>"
        f"<table style='border-collapse:collapse'>{rows}</table>"
        f"{table_html}"
    )


def _fill_docx_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """Подстановка {{key}} в абзаце. python-docx дробит текст по runs — если в абзаце
    есть плейсхолдер, переписываем текст в первый run (вёрстка поля простая — приемлемо)."""
    text = paragraph.text
    if "{{" not in text:
        return
    for key, val in mapping.items():
        text = text.replace("{{" + key + "}}", val)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def render_docx(resolved: dict[str, Any], out_path: Path, template_path: Path | None = None) -> Path:
    from docx import Document

    if template_path and Path(template_path).is_file():
        doc = Document(str(template_path))
        mapping = _value_map(resolved)
        for p in doc.paragraphs:
            _fill_docx_paragraph(p, mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _fill_docx_paragraph(p, mapping)
    else:
        # template-less fallback: документ строится из дескриптора
        doc = Document()
        doc.add_heading(resolved["title"], level=1)
        if resolved.get("legal_basis"):
            doc.add_paragraph(resolved["legal_basis"]).italic = True
        table = doc.add_table(rows=0, cols=2)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for f in resolved["fields"]:
            cells = table.add_row().cells
            cells[0].text = f["label"]
            cells[1].text = f["value"] or "—"
    cols = resolved.get("columns") or []
    if cols:
        doc.add_paragraph()
        gt = doc.add_table(rows=1, cols=len(cols))
        try:
            gt.style = "Table Grid"
        except Exception:
            pass
        for i, c in enumerate(cols):
            gt.rows[0].cells[i].text = str(c)
        for r in resolved.get("rows", []):
            cells = gt.add_row().cells
            for i, val in enumerate(r):
                cells[i].text = str(val)
    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path


def render_xlsx(resolved: dict[str, Any], out_path: Path) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = resolved["id"][:31]
    ws["A1"] = resolved["title"]
    ws["A1"].font = Font(bold=True, size=13)
    ws["A2"] = resolved.get("legal_basis", "")
    ws["A2"].font = Font(italic=True, size=9)
    r = 4
    for f in resolved["fields"]:
        ws.cell(r, 1, f["label"]).font = Font(color="555555")
        ws.cell(r, 2, f["value"]).font = Font(bold=True)
        r += 1
    ws.column_dimensions["A"].width = 48
    ws.column_dimensions["B"].width = 40

    cols = resolved.get("columns") or []
    if cols:
        r += 1
        for c, title in enumerate(cols, 1):
            cell = ws.cell(r, c, title)
            cell.font = Font(bold=True)
        for row_vals in resolved.get("rows", []):
            r += 1
            for c, val in enumerate(row_vals, 1):
                ws.cell(r, c, val)
        for c in range(1, len(cols) + 1):
            ws.column_dimensions[chr(64 + c) if c <= 26 else "AA"].width = 22

    out_path = Path(out_path)
    wb.save(str(out_path))
    return out_path


def generate(
    form_id: str, fmt: str, *, project_id: int | None = None, manual: dict[str, Any] | None = None,
    out_path: Path | None = None,
) -> dict[str, Any]:
    """Сгенерировать документ формы в формате fmt. Возвращает путь и резолв полей."""
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"fmt: {list(SUPPORTED_FORMATS)}")
    resolved = resolve_fields(form_id, project_id, manual)
    if resolved is None:
        raise ValueError(f"Форма {form_id!r} не найдена")
    descriptor = load_descriptor(form_id) or {}

    if fmt == "html":
        return {"resolved": resolved, "html": render_html(resolved), "path": None}

    if out_path is None:
        stamp = time.strftime("%Y%m%d_%H%M%S", time.localtime())
        out_path = _output_dir() / f"{form_id}_{project_id or 0}_{stamp}.{fmt}"
    out_path = Path(out_path)

    if fmt == "docx":
        tmpl = (descriptor.get("templates") or {}).get("docx")
        tmpl_path = Path(tmpl) if tmpl else None
        render_docx(resolved, out_path, tmpl_path)
    elif fmt == "xlsx":
        render_xlsx(resolved, out_path)
    return {"resolved": resolved, "html": None, "path": str(out_path)}
