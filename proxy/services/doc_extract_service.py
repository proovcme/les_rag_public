"""Unified Construction Harness v0.13 — Document Body Extraction / Ingestion-Lite (read-only).

Готовит ТЕКСТОВЫЙ слой из PDF/DOCX/XLSX в sidecar-JSONL рядом с датасетом (НЕ трогая оригиналы,
БЕЗ OCR, без облака). source_ref до page/paragraph/sheet/row. Нет библиотеки → unavailable; PDF без
text-слоя → no_text_layer (actionable, не фейк). Sidecar: storage/datasets/{ds}/_extracted/<rel>.jsonl.

extractor_version фиксируется в каждом item; пустой/битый extraction → status, не молчаливый пропуск.
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

EXTRACTOR_VERSION = "v0.13"
SIDECAR_DIRNAME = "_extracted"
_SUPPORTED = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".csv"}
_MAX_BYTES_DEFAULT = 40 * 1024 * 1024


@dataclass
class SidecarItem:
    dataset_id: str
    original_file_name: str
    original_relative_path: str
    source_kind: str                      # pdf_text|docx_text|docx_table|xlsx_row|txt_body|md_body|csv_row
    text: str
    source_ref: str
    page: int | None = None
    paragraph_index: int | None = None
    row_index: int | None = None
    sheet_name: str | None = None
    extractor_version: str = EXTRACTOR_VERSION
    extraction_status: str = "ok"         # ok|no_text_layer|failed|unavailable
    warnings: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {k: v for k, v in self.__dict__.items()}


@dataclass
class ExtractResult:
    status: str                           # ok|no_text_layer|unavailable|failed|skipped
    source_kind: str = ""
    items: list[SidecarItem] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _ref(ds: str, rel: str, suffix: str) -> str:
    return f"{ds}/{rel}#{suffix}"


# ── PDF (fitz/PyMuPDF; no_text_layer без OCR) ─────────────────────────────────────────────

def extract_pdf_text(path: Path, *, ds: str, rel: str) -> ExtractResult:
    try:
        import fitz  # PyMuPDF
    except Exception:
        try:
            import pdfplumber  # noqa: F401
            return _extract_pdf_pdfplumber(path, ds=ds, rel=rel)
        except Exception:
            return ExtractResult("unavailable", "pdf_text",
                                 warnings=["pdf_text_extractor_unavailable: нет fitz/pdfplumber"])
    try:
        doc = fitz.open(path)
    except Exception as e:  # noqa: BLE001
        return ExtractResult("failed", "pdf_text", warnings=[f"pdf_open_failed: {str(e)[:60]}"])
    items, total_chars = [], 0
    for i in range(doc.page_count):
        try:
            text = doc.load_page(i).get_text("text") or ""
        except Exception:  # noqa: BLE001
            text = ""
        text = text.strip()
        total_chars += len(text)
        if text:
            items.append(SidecarItem(ds, path.name, rel, "pdf_text", text, _ref(ds, rel, f"p{i+1}"), page=i + 1))
    doc.close()
    if total_chars == 0:
        return ExtractResult("no_text_layer", "pdf_text", warnings=[
            "no_text_layer: PDF без текстового слоя — нужен OCR/ingestion вне hot-path"])
    return ExtractResult("ok", "pdf_text", items=items)


def _extract_pdf_pdfplumber(path: Path, *, ds: str, rel: str) -> ExtractResult:
    import pdfplumber
    items, total = [], 0
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()
            total += len(text)
            if text:
                items.append(SidecarItem(ds, path.name, rel, "pdf_text", text, _ref(ds, rel, f"p{i+1}"), page=i + 1))
    if total == 0:
        return ExtractResult("no_text_layer", "pdf_text", warnings=["no_text_layer: PDF без текстового слоя"])
    return ExtractResult("ok", "pdf_text", items=items)


# ── DOCX (python-docx) ───────────────────────────────────────────────────────────────────

def extract_docx(path: Path, *, ds: str, rel: str) -> ExtractResult:
    try:
        from docx import Document
    except Exception:
        return ExtractResult("unavailable", "docx_text", warnings=["docx_extractor_unavailable: нет python-docx"])
    try:
        d = Document(str(path))
    except Exception as e:  # noqa: BLE001
        return ExtractResult("failed", "docx_text", warnings=[f"docx_open_failed: {str(e)[:60]}"])
    items = []
    for i, p in enumerate(d.paragraphs):
        t = (p.text or "").strip()
        if t:
            items.append(SidecarItem(ds, path.name, rel, "docx_text", t, _ref(ds, rel, f"para{i}"), paragraph_index=i))
    for ti, tbl in enumerate(d.tables):
        for ri, row in enumerate(tbl.rows):
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                items.append(SidecarItem(ds, path.name, rel, "docx_table", line,
                                         _ref(ds, rel, f"t{ti}r{ri}"), row_index=ri))
    if not items:
        return ExtractResult("no_text_layer", "docx_text", warnings=["docx пуст или без текста"])
    return ExtractResult("ok", "docx_text", items=items)


# ── XLSX generic (openpyxl; НЕ resource-workbook, общий табличный) ───────────────────────

def extract_xlsx_generic(path: Path, *, ds: str, rel: str, max_rows: int = 5000) -> ExtractResult:
    try:
        import openpyxl
    except Exception:
        return ExtractResult("unavailable", "xlsx_row", warnings=["xlsx_extractor_unavailable: нет openpyxl"])
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception as e:  # noqa: BLE001
        return ExtractResult("failed", "xlsx_row", warnings=[f"xlsx_open_failed: {str(e)[:60]}"])
    items, seen = [], 0
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for ri, row in enumerate(ws.iter_rows(values_only=True), 1):
            if seen >= max_rows:
                break
            vals = [str(v) for v in row if v not in (None, "")]
            if not vals:
                continue
            seen += 1
            line = " | ".join(vals)
            items.append(SidecarItem(ds, path.name, rel, "xlsx_row", line,
                                     f"{ds}/{rel}#{sheet}!R{ri}", row_index=ri, sheet_name=sheet))
    wb.close()
    if not items:
        return ExtractResult("no_text_layer", "xlsx_row", warnings=["xlsx без непустых строк"])
    return ExtractResult("ok", "xlsx_row", items=items)


# ── txt/md/csv (тривиально — для полноты sidecar; .md/.txt также читает file_body напрямую) ─

def extract_text_file(path: Path, *, ds: str, rel: str, kind: str) -> ExtractResult:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:  # noqa: BLE001
        return ExtractResult("failed", kind, warnings=[f"read_failed: {str(e)[:60]}"])
    items = []
    for i, ln in enumerate(text.splitlines(), 1):
        if ln.strip():
            items.append(SidecarItem(ds, path.name, rel, kind, ln.strip(), _ref(ds, rel, f"L{i}"), page=None))
    return ExtractResult("ok" if items else "no_text_layer", kind, items=items)


# ── dispatcher ────────────────────────────────────────────────────────────────────────────

def extract_file(path: Path, *, ds: str, rel: str) -> ExtractResult:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_text(path, ds=ds, rel=rel)
    if ext == ".docx":
        return extract_docx(path, ds=ds, rel=rel)
    if ext == ".xlsx":
        return extract_xlsx_generic(path, ds=ds, rel=rel)
    if ext in (".txt",):
        return extract_text_file(path, ds=ds, rel=rel, kind="txt_body")
    if ext in (".md",):
        return extract_text_file(path, ds=ds, rel=rel, kind="md_body")
    if ext in (".csv",):
        return extract_text_file(path, ds=ds, rel=rel, kind="csv_row")
    return ExtractResult("skipped", warnings=[f"unsupported_binary: {ext}"])


# ── sidecar I/O ───────────────────────────────────────────────────────────────────────────

def sidecar_path(storage_root: Path, ds: str, rel: str) -> Path:
    return storage_root / ds / SIDECAR_DIRNAME / (rel + ".jsonl")


def write_sidecar(storage_root: Path, ds: str, rel: str, items: list[SidecarItem]) -> Path:
    sp = sidecar_path(storage_root, ds, rel)
    sp.parent.mkdir(parents=True, exist_ok=True)
    with sp.open("w", encoding="utf-8") as f:
        for it in items:
            f.write(json.dumps(it.as_dict(), ensure_ascii=False) + "\n")
    return sp


def read_sidecars(storage_root: Path, ds: str) -> list[dict[str, Any]]:
    """Прочитать все sidecar-items датасета (для search_extracted_body / index_health)."""
    base = storage_root / ds / SIDECAR_DIRNAME
    out: list[dict] = []
    if not base.exists():
        return out
    for p in base.rglob("*.jsonl"):
        try:
            for ln in p.read_text(encoding="utf-8").splitlines():
                if ln.strip():
                    out.append(json.loads(ln))
        except Exception:  # noqa: BLE001
            continue
    return out


def sidecar_count(storage_root: Path, ds: str) -> int:
    base = storage_root / ds / SIDECAR_DIRNAME
    return sum(1 for _ in base.rglob("*.jsonl")) if base.exists() else 0


# ── ВОР-таблицы из XLSX/DOCX (header+rows) — для bor_extract/ЛСР без parquet ──────────────

def extract_bor_tables(path: Path) -> list[dict[str, Any]]:
    """Вернуть простые таблицы {header:[...], rows:[[...]], line_start} из .xlsx/.docx. Маппинг в
    name/unit/qty делает markdown_table_to_rows (общие синонимы). Без оценки формул — только значения."""
    ext = path.suffix.lower()
    out: list[dict] = []
    if ext == ".xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        except Exception:  # noqa: BLE001
            return out
        for sheet in wb.sheetnames:
            rows = [[("" if v is None else str(v)) for v in r]
                    for r in wb[sheet].iter_rows(values_only=True)]
            rows = [r for r in rows if any(c.strip() for c in r)]
            if len(rows) >= 2:
                out.append({"header": rows[0], "rows": rows[1:], "line_start": 1, "line_end": len(rows),
                            "sheet": sheet})
        wb.close()
    elif ext == ".docx":
        try:
            from docx import Document
            d = Document(str(path))
        except Exception:  # noqa: BLE001
            return out
        for tbl in d.tables:
            rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
            rows = [r for r in rows if any(r)]
            if len(rows) >= 2:
                out.append({"header": rows[0], "rows": rows[1:], "line_start": 1, "line_end": len(rows)})
    return out
