"""Office and Structured Document Passport Service for Л.И.С.Т. / LES.

Produces structured passports (list.office_passport.v1) for Excel, Word, PowerPoint, EML,
and tabular text files. Complements PDF contour passports for full-corpus indexing.
"""
from __future__ import annotations

import csv
import io
import re
import zipfile
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

OFFICE_PASSPORT_SCHEMA = "list.office_passport.v1"


def audit_office_document(path: str | Path) -> dict[str, Any]:
    """Generate a structured passport for non-PDF document files."""
    file_path = Path(path).expanduser().resolve()
    if not file_path.is_file():
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    suffix = file_path.suffix.lower()
    passport: dict[str, Any] = {
        "schema": OFFICE_PASSPORT_SCHEMA,
        "file_name": file_path.name,
        "suffix": suffix,
        "size_bytes": file_path.stat().st_size,
        "passport_kind": "unknown",
        "structure": {},
        "warnings": [],
    }

    try:
        if suffix in {".xlsx", ".xlsm"}:
            passport["passport_kind"] = "excel"
            passport["structure"] = _audit_excel(file_path)
        elif suffix == ".docx":
            passport["passport_kind"] = "word"
            passport["structure"] = _audit_word(file_path)
        elif suffix == ".pptx":
            passport["passport_kind"] = "presentation"
            passport["structure"] = _audit_pptx(file_path)
        elif suffix == ".eml":
            passport["passport_kind"] = "mail"
            passport["structure"] = _audit_eml(file_path)
        elif suffix in {".csv", ".tsv"}:
            passport["passport_kind"] = "table"
            passport["structure"] = _audit_csv(file_path)
        elif suffix in {".txt", ".md", ".json", ".xml", ".yaml", ".yml"}:
            passport["passport_kind"] = "text"
            passport["structure"] = _audit_text(file_path)
        else:
            passport["warnings"].append("Универсальный просмотрщик / неструктурированный документ")
    except Exception as exc:  # noqa: BLE001
        passport["warnings"].append(f"Ошибка паспортизации: {type(exc).__name__}: {exc}")

    return passport


def _audit_excel(path: Path) -> dict[str, Any]:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=False)
    try:
        sheets_info = []
        total_rows = 0
        total_cols = 0
        has_any_formulas = False

        for name in wb.sheetnames:
            ws = wb[name]
            m_row = ws.max_row or 0
            m_col = ws.max_column or 0
            total_rows += m_row
            total_cols = max(total_cols, m_col)
            sheets_info.append({
                "name": name,
                "max_row": m_row,
                "max_column": m_col,
            })

        return {
            "sheet_count": len(wb.sheetnames),
            "sheet_names": list(wb.sheetnames),
            "sheets": sheets_info,
            "total_rows": total_rows,
            "max_columns": total_cols,
        }
    finally:
        wb.close()


def _audit_word(path: Path) -> dict[str, Any]:
    from docx import Document

    doc = Document(str(path))
    headings: list[dict[str, Any]] = []
    paragraph_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    for idx, p in enumerate(doc.paragraphs[:1000]):
        text = (p.text or "").strip()
        style = str(getattr(p.style, "name", "") or "").lower()
        if text and ("heading" in style or "заголовок" in style or "title" in style):
            headings.append({
                "anchor": f"para{idx}",
                "text": text[:120],
                "style": style,
            })

    return {
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "heading_count": len(headings),
        "headings": headings[:30],
    }


def _audit_pptx(path: Path) -> dict[str, Any]:
    slide_count = 0
    with zipfile.ZipFile(path) as archive:
        slide_count = len([name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)])
    return {
        "slide_count": slide_count,
    }


def _audit_eml(path: Path) -> dict[str, Any]:
    msg = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    attachments = []
    if msg.is_multipart():
        for part in msg.walk():
            filename = part.get_filename()
            if filename:
                attachments.append(filename)
    return {
        "from": str(msg.get("From", "")),
        "to": str(msg.get("To", "")),
        "subject": str(msg.get("Subject", "")),
        "date": str(msg.get("Date", "")),
        "attachment_count": len(attachments),
        "attachments": attachments,
    }


def _audit_csv(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")[:64000]
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        delimiter = "\t" if path.suffix.lower() == ".tsv" else ","

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = list(reader)
    col_count = max((len(r) for r in rows), default=0)
    return {
        "row_count": len(rows),
        "column_count": col_count,
        "delimiter": delimiter,
    }


def _audit_text(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    return {
        "line_count": len(lines),
        "char_count": len(text),
    }
