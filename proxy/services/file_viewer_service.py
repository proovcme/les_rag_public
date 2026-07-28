"""Read-only, dependency-local previews for evidence files in the LES GUI.

This is intentionally a viewer, not an editor or a fidelity converter.  It
opens originals only for reading and produces escaped HTML for the browser and
Tauri/WebView2 shells.  PDF rendering remains in ``pdf_viewer_service`` because
it has page/evidence-coordinate semantics of its own.
"""
from __future__ import annotations

import csv
import html
import io
import json
import re
import zipfile
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any
from xml.etree import ElementTree


FILE_VIEWER_SCHEMA = "list.file_viewer.v1"
MAX_TEXT_CHARS = 600_000
MAX_TABLE_ROWS = 500
MAX_TABLE_COLUMNS = 80

_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
_TEXT_SUFFIXES = {
    ".txt", ".md", ".json", ".jsonl", ".xml", ".yaml", ".yml", ".log",
    ".ini", ".cfg", ".sql", ".py", ".html", ".svg",
}
_TABLE_SUFFIXES = {".csv", ".tsv"}
VIEWABLE_SUFFIXES = {
    ".pdf", ".docx", ".xlsx", ".xlsm", ".pptx", ".eml",
    *_IMAGE_SUFFIXES, *_TEXT_SUFFIXES, *_TABLE_SUFFIXES,
}


def is_viewable_file(path: str | Path) -> bool:
    return Path(path).suffix.lower() in VIEWABLE_SUFFIXES


def viewer_kind(path: str | Path) -> str:
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "word"
    if suffix in {".xlsx", ".xlsm"}:
        return "excel"
    if suffix == ".pptx":
        return "presentation"
    if suffix == ".eml":
        return "mail"
    if suffix in _IMAGE_SUFFIXES:
        return "image"
    if suffix in _TABLE_SUFFIXES:
        return "table"
    if suffix in _TEXT_SUFFIXES:
        return "text"
    return "unsupported"


def _e(value: Any) -> str:
    return html.escape(str(value if value is not None else ""), quote=True)


def _safe_json(value: object) -> str:
    return json.dumps(value, ensure_ascii=False, separators=(",", ":")).replace("<", "\\u003c")


def _document_shell(*, path_id: str, file_name: str, kind: str, body: str, note: str = "") -> str:
    config = _safe_json({
        "schema": FILE_VIEWER_SCHEMA,
        "path": str(path_id),
        "name": str(file_name),
        "kind": kind,
    })
    note_html = f'<div class="notice">{_e(note)}</div>' if note else ""
    template = """<!doctype html>
<html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Л.И.С.Т. — __TITLE__</title>
<style>
:root{color-scheme:light;--bg:#eef3f7;--panel:#fff;--text:#18242e;--dim:#647482;--accent:#0f8b68;
--line:rgba(24,36,46,.13);--soft:rgba(15,139,104,.08);--shadow:0 12px 34px rgba(20,36,50,.12)}
*{box-sizing:border-box}html,body{margin:0;min-height:100%;background:var(--bg);color:var(--text);
font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif;-webkit-font-smoothing:antialiased}
body{min-height:100vh}.toolbar{position:sticky;top:0;z-index:8;display:flex;align-items:center;gap:8px;
min-height:52px;padding:6px 12px;background:rgba(255,255,255,.95);box-shadow:0 1px 0 var(--line),0 8px 24px rgba(20,36,50,.07);backdrop-filter:blur(12px)}
.name{min-width:0;flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;font-size:12px;font-weight:800}
.kind{padding:4px 8px;border-radius:7px;background:var(--soft);color:var(--accent);font-size:10px;font-weight:850;text-transform:uppercase;letter-spacing:.04em}
button,.open{display:inline-flex;align-items:center;justify-content:center;min-height:40px;padding:0 12px;border:0;border-radius:9px;background:transparent;color:var(--text);font:750 12px/1 inherit;text-decoration:none;cursor:pointer;transition-property:background-color,color,scale;transition-duration:.14s}
button:hover,.open:hover{color:var(--accent);background:var(--soft)}button:active,.open:active{scale:.96}
.content{width:min(1180px,calc(100% - 28px));margin:18px auto 40px}.notice{margin-bottom:12px;padding:9px 11px;border:1px solid var(--line);border-radius:10px;background:#fff9e8;color:#765b20;font-size:12px;line-height:1.45}
.document{width:min(860px,100%);margin:auto;padding:52px 64px;background:var(--panel);outline:1px solid rgba(0,0,0,.08);box-shadow:var(--shadow)}
.document p{margin:0 0 10px;line-height:1.55;white-space:pre-wrap}.document h1,.document h2,.document h3{margin:1.25em 0 .55em;line-height:1.25}.document h1{font-size:24px}.document h2{font-size:19px}.document h3{font-size:15px}
.anchor{scroll-margin-top:72px}.evidence-hit{border-radius:5px;background:rgba(255,215,64,.28);box-shadow:0 0 0 4px rgba(255,215,64,.15)}
.table-wrap{max-width:100%;margin:14px 0;overflow:auto;border:1px solid var(--line);border-radius:10px;background:var(--panel);box-shadow:0 4px 14px rgba(20,36,50,.05)}
table{border-collapse:separate;border-spacing:0;width:100%;font-size:12px;font-variant-numeric:tabular-nums}th,td{min-width:72px;padding:7px 9px;border-right:1px solid var(--line);border-bottom:1px solid var(--line);text-align:left;vertical-align:top;white-space:pre-wrap}th{position:sticky;top:0;z-index:2;background:#f5f8fa;color:var(--dim);font-weight:800}tr:last-child td{border-bottom:0}th:last-child,td:last-child{border-right:0}.row-number{min-width:48px;width:48px;color:var(--dim);text-align:right;background:#f7f9fb;position:sticky;left:0;z-index:1}.formula{color:#705b20;font-family:ui-monospace,SFMono-Regular,Menlo,monospace}
.sheets{display:flex;gap:6px;overflow:auto;margin:0 0 10px;padding:2px}.sheet{white-space:nowrap;background:var(--panel);box-shadow:inset 0 0 0 1px var(--line)}.sheet.active{color:#fff;background:var(--accent);box-shadow:none}
.text-view{margin:0;padding:18px;overflow:auto;border-radius:10px;background:#101820;color:#e8f0f5;box-shadow:var(--shadow);font:12px/1.6 ui-monospace,SFMono-Regular,Menlo,monospace;white-space:pre-wrap;overflow-wrap:anywhere}
.image-stage{display:grid;place-items:center;min-height:calc(100vh - 100px)}.image-stage img{display:block;max-width:100%;max-height:calc(100vh - 110px);background:#fff;outline:1px solid rgba(0,0,0,.12);box-shadow:var(--shadow)}
.slides{display:grid;gap:18px}.slide{aspect-ratio:16/9;padding:42px 52px;border-radius:4px;background:#fff;outline:1px solid rgba(0,0,0,.1);box-shadow:var(--shadow);overflow:auto}.slide-number{margin-bottom:16px;color:var(--dim);font-size:11px;font-weight:800}.slide p{font-size:16px;line-height:1.45}.mail-card{padding:26px;border-radius:4px;background:#fff;outline:1px solid rgba(0,0,0,.1);box-shadow:var(--shadow)}.mail-head{display:grid;grid-template-columns:auto 1fr;gap:7px 14px;padding-bottom:18px;border-bottom:1px solid var(--line);font-size:12px}.mail-key{color:var(--dim);font-weight:800}.mail-body{margin-top:20px;line-height:1.6;white-space:pre-wrap}
.empty{padding:36px;border:1px dashed var(--line);border-radius:12px;background:var(--panel);color:var(--dim);text-align:center;line-height:1.5}
@media(max-width:700px){.content{width:calc(100% - 16px);margin-top:10px}.document{padding:28px 22px}.kind{display:none}.slide{padding:26px 28px}.toolbar{padding-inline:8px}}
</style></head><body>
<header class="toolbar"><div class="name" id="name"></div><span class="kind">__KIND__</span><a class="open" id="open" target="_blank" rel="noopener">Оригинал</a></header>
<main class="content">__NOTE____BODY__</main>
<script id="viewer-config" type="application/json">__CONFIG__</script>
<script>
const cfg=JSON.parse(document.getElementById('viewer-config').textContent);const api=location.pathname.startsWith('/lite-api/')?'/lite-api':'/api';
document.getElementById('name').textContent=cfg.name;const raw=new URLSearchParams({path:cfg.path});document.getElementById('open').href=`${api}/rag/file/raw?${raw.toString()}`;
const rawImage=document.querySelector('[data-raw-image]');if(rawImage)rawImage.src=`${api}/rag/file/raw?${raw.toString()}`;
document.querySelectorAll('[data-sheet]').forEach(button=>button.addEventListener('click',()=>{const url=new URL(location.href);url.searchParams.set('sheet',button.dataset.sheet);url.searchParams.delete('locator');location.href=url.toString()}));
const hit=document.querySelector('.evidence-hit');if(hit)setTimeout(()=>hit.scrollIntoView({block:'center'}),40);
</script></body></html>"""
    return (template.replace("__TITLE__", _e(file_name))
            .replace("__KIND__", _e(kind))
            .replace("__NOTE__", note_html)
            .replace("__BODY__", body)
            .replace("__CONFIG__", config))


def _docx_view(path: Path, locator: str) -> tuple[str, str]:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    paragraphs = document.paragraphs[:2000]

    def render_paragraph(paragraph, index: int) -> str:
        value = paragraph.text or ""
        style = str(getattr(paragraph.style, "name", "") or "").lower()
        tag = "p"
        if "heading 1" in style or "заголовок 1" in style:
            tag = "h1"
        elif "heading 2" in style or "заголовок 2" in style:
            tag = "h2"
        elif "heading" in style or "заголовок" in style or "title" in style:
            tag = "h3"
        anchor = f"para{index}"
        hit = " evidence-hit" if locator.lower().startswith(anchor.lower()) else ""
        return f'<{tag} id="{anchor}" class="anchor{hit}">{_e(value) or "&nbsp;"}</{tag}>'

    def render_table(table, table_index: int) -> str:
        rows: list[str] = []
        for row_index, row in enumerate(table.rows[:MAX_TABLE_ROWS]):
            cells = "".join(f"<td>{_e(cell.text)}</td>" for cell in row.cells[:MAX_TABLE_COLUMNS])
            rows.append(f'<tr id="t{table_index}r{row_index}">{cells}</tr>')
        return f'<div class="table-wrap anchor" id="t{table_index}"><table>{"".join(rows)}</table></div>'

    # python-docx 1.2 exposes body items in document order.  Keeping that order
    # matters for review even though this preview intentionally does not mimic
    # Word's pagination.
    paragraph_index = 0
    table_index = 0
    for block in document.iter_inner_content():
        if getattr(block, "_p", None) is not None:
            if paragraph_index < len(paragraphs):
                parts.append(render_paragraph(block, paragraph_index))
            paragraph_index += 1
            continue
        if getattr(block, "_tbl", None) is not None and table_index < 150:
            parts.append(render_table(block, table_index))
            table_index += 1
    note = ""
    if len(document.paragraphs) > len(paragraphs):
        note = f"Показаны первые {len(paragraphs)} абзацев. Для полной проверки откройте оригинал."
    body = f'<article class="document">{"".join(parts) or "<div class=\"empty\">Документ не содержит отображаемого текста.</div>"}</article>'
    return body, note


def _xlsx_view(path: Path, *, sheet: str, locator: str) -> tuple[str, str]:
    import openpyxl
    from openpyxl.utils import get_column_letter

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=False)
    try:
        loc_match = re.fullmatch(r"(.+)!R(\d+)(?:C\d+)?", locator or "", re.I)
        locator_sheet = loc_match.group(1) if loc_match else ""
        target_row = int(loc_match.group(2)) if loc_match else 0
        selected = sheet if sheet in workbook.sheetnames else locator_sheet
        if selected not in workbook.sheetnames:
            selected = workbook.sheetnames[0]
        worksheet = workbook[selected]
        start_row = max(1, target_row - 40) if target_row else 1
        end_row = start_row + MAX_TABLE_ROWS - 1
        rows_data: list[list[Any]] = []
        max_columns = 1
        for row in worksheet.iter_rows(min_row=start_row, max_row=end_row, values_only=True):
            values = list(row[:MAX_TABLE_COLUMNS])
            while values and values[-1] is None:
                values.pop()
            rows_data.append(values)
            max_columns = max(max_columns, len(values))
        header = "".join(f"<th>{get_column_letter(index)}</th>" for index in range(1, max_columns + 1))
        rendered_rows: list[str] = []
        for offset, values in enumerate(rows_data):
            row_number = start_row + offset
            classes = " evidence-hit" if row_number == target_row and selected == locator_sheet else ""
            cells: list[str] = []
            for value in values:
                cell_class = ' class="formula"' if isinstance(value, str) and value.startswith("=") else ""
                cells.append(f"<td{cell_class}>{_e(value)}</td>")
            cells.extend("<td></td>" for _ in range(max_columns - len(values)))
            rendered_rows.append(
                f'<tr id="row-{row_number}" class="anchor{classes}"><th class="row-number">{row_number}</th>{"".join(cells)}</tr>'
            )
        tabs = "".join(
            f'<button class="sheet{" active" if name == selected else ""}" data-sheet="{_e(name)}">{_e(name)}</button>'
            for name in workbook.sheetnames
        )
        table = f'<div class="sheets">{tabs}</div><div class="table-wrap"><table><thead><tr><th class="row-number"></th>{header}</tr></thead><tbody>{"".join(rendered_rows)}</tbody></table></div>'
        note_parts = [f"Лист: {selected}. Показаны строки {start_row}–{start_row + max(0, len(rows_data) - 1)}."]
        if worksheet.max_row > end_row or worksheet.max_column > MAX_TABLE_COLUMNS:
            note_parts.append("Большая таблица ограничена окном просмотра; оригинал доступен кнопкой сверху.")
        return table, " ".join(note_parts)
    finally:
        workbook.close()


def _delimited_view(path: Path) -> tuple[str, str]:
    text = path.read_text(encoding="utf-8", errors="replace")[:MAX_TEXT_CHARS]
    sample = text[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel_tab if path.suffix.lower() == ".tsv" else csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect=dialect))[:MAX_TABLE_ROWS]
    width = min(MAX_TABLE_COLUMNS, max((len(row) for row in rows), default=1))
    rendered = []
    for index, row in enumerate(rows, 1):
        cells = "".join(f"<td>{_e(value)}</td>" for value in row[:width])
        rendered.append(f'<tr><th class="row-number">{index}</th>{cells}</tr>')
    return f'<div class="table-wrap"><table><tbody>{"".join(rendered)}</tbody></table></div>', (
        f"Показаны первые {len(rows)} строк." if len(rows) >= MAX_TABLE_ROWS else ""
    )


def _pptx_view(path: Path, locator: str) -> tuple[str, str]:
    slides: list[tuple[int, str]] = []
    with zipfile.ZipFile(path) as archive:
        names = sorted(
            (name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)),
            key=lambda name: int(re.search(r"(\d+)", Path(name).stem).group(1)),
        )
        for name in names[:300]:
            root = ElementTree.fromstring(archive.read(name))
            text_nodes = [node.text or "" for node in root.iter() if node.tag.endswith("}t") and (node.text or "").strip()]
            number = int(re.search(r"(\d+)", Path(name).stem).group(1))
            slides.append((number, "\n".join(text_nodes)))
    match = re.search(r"(?:slide|слайд)(\d+)", locator or "", re.I)
    target = int(match.group(1)) if match else 0
    cards = "".join(
        f'<section id="slide{number}" class="slide anchor{" evidence-hit" if number == target else ""}"><div class="slide-number">Слайд {number}</div><p>{_e(text)}</p></section>'
        for number, text in slides
    )
    return f'<div class="slides">{cards or "<div class=\"empty\">На слайдах нет извлекаемого текста.</div>"}</div>', "Предпросмотр показывает текст и порядок слайдов без изменения презентации."


def _mail_view(path: Path) -> tuple[str, str]:
    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    fields = [("От", message.get("From", "")), ("Кому", message.get("To", "")),
              ("Дата", message.get("Date", "")), ("Тема", message.get("Subject", ""))]
    body = message.get_body(preferencelist=("plain",)) if message.is_multipart() else message
    text = ""
    if body is not None:
        try:
            text = body.get_content()
        except (LookupError, UnicodeError):
            text = str(body.get_payload(decode=True) or b"", errors="replace")
    head = "".join(f'<div class="mail-key">{_e(key)}</div><div>{_e(value)}</div>' for key, value in fields)
    return f'<article class="mail-card"><div class="mail-head">{head}</div><div class="mail-body">{_e(text[:MAX_TEXT_CHARS])}</div></article>', "Вложения письма открываются как самостоятельные источники."


def file_viewer_html(
    source_path: str | Path,
    *,
    path_id: str,
    locator: str = "",
    sheet: str = "",
) -> str:
    path = Path(source_path).expanduser().resolve()
    if not path.is_file():
        raise FileNotFoundError("Файл не найден")
    kind = viewer_kind(path)
    if kind == "unsupported":
        raise ValueError("Встроенный просмотр для этого формата пока недоступен")
    try:
        if kind == "word":
            body, note = _docx_view(path, locator)
        elif kind == "excel":
            body, note = _xlsx_view(path, sheet=sheet, locator=locator)
        elif kind == "presentation":
            body, note = _pptx_view(path, locator)
        elif kind == "mail":
            body, note = _mail_view(path)
        elif kind == "table":
            body, note = _delimited_view(path)
        elif kind == "image":
            body, note = '<div class="image-stage"><img data-raw-image alt="Изображение документа"></div>', ""
        elif kind == "text":
            value = path.read_text(encoding="utf-8", errors="replace")
            clipped = value[:MAX_TEXT_CHARS]
            body = f'<pre class="text-view">{_e(clipped)}</pre>'
            note = "Файл сокращён в предпросмотре." if len(value) > len(clipped) else ""
        else:
            raise ValueError("PDF использует отдельный постраничный просмотрщик")
    except (zipfile.BadZipFile, ElementTree.ParseError, KeyError) as exc:
        raise ValueError("Файл повреждён или имеет несовместимую структуру") from exc
    return _document_shell(path_id=path_id, file_name=path.name, kind=kind, body=body, note=note)
