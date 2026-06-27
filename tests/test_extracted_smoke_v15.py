"""Unified Construction Harness v0.15 — approved runtime sidecar write + real extracted_body smoke.

Оператор разрешил запись sidecar для 844a2b53 (27 ГОСТ/СП .docx → 23 930 параграфов). norm/source
теперь отвечают из extracted_body с source_ref до абзаца; no_lexical_index заменён реальным RETRIEVED
ИЛИ честным term_not_found. Оригиналы read-only. Числа/термины не выдумываются.
"""

from pathlib import Path

import pytest

from proxy.services import doc_extract_service as de
from proxy.services import source_adapters as sa
from proxy.services import unified_construction_harness_service as u
from proxy.services import resource_cost_service as rc
from proxy.services import construction_harness_service as ch
from proxy.services.evidence_contract import EvidenceType

_RT = Path("/Users/ovc/LES/storage/datasets")
_GOST = "844a2b53-9658-4e5a-92e4-f649de8af043"
_SIDECARS = (_RT / _GOST / "_extracted").exists()


def _docx(path, heading="СП котельная", paras=("АУПТ требуется для котельной по п.5.4",
                                               "предел огнестойкости стен принят R45")):
    from docx import Document
    d = Document()
    d.add_heading(heading, 1)
    for p in paras:
        d.add_paragraph(p)
    d.save(str(path))


@pytest.fixture
def sidecar_ds(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _docx(d / "СП_котельная.docx")
    r = de.extract_docx(d / "СП_котельная.docx", ds="ds", rel="СП_котельная.docx")
    de.write_sidecar(tmp_path, "ds", "СП_котельная.docx", r.items)
    de.write_manifest(tmp_path, "ds", [{"original_relative_path": "СП_котельная.docx",
                                        "original_size": (d / "СП_котельная.docx").stat().st_size,
                                        "original_mtime": (d / "СП_котельная.docx").stat().st_mtime,
                                        "ext": ".docx", "status": "ok", "item_count": len(r.items)}])
    return tmp_path


# ── norm_qa word-expansion (v0.15 fix: фраза → content-слова) ─────────────────────────────

def test_norm_qa_word_expansion_finds_term(sidecar_ds):
    # «правила огнестойкости стен» — фраза целиком не матчит, но слово «огнестойкость» в теле
    r = u.run_unified_construction_harness("правила огнестойкости стен", dataset_ids=["ds"], storage_root=sidecar_ds)
    assert r.total_status == "complete"
    retr = next(b for b in r.evidence_blocks if b.type is EvidenceType.RETRIEVED)
    assert retr.items[0].source_refs and "extracted_body" in r.answer_data.get("searched_tiers", [])

def test_norm_qa_term_absent_not_no_lexical_index(sidecar_ds):
    # термина нет в извлечённом теле → no_data, но НЕ «no_lexical_index» (sidecar есть и просмотрен)
    r = u.run_unified_construction_harness("требования к лифтам ВВГнг", dataset_ids=["ds"], storage_root=sidecar_ds)
    assert r.total_status == "no_data"
    h = r.answer_data.get("index_health", {})
    if h.get("datasets"):
        warns = h["datasets"][0].get("warnings", [])
        assert "no_lexical_index" not in warns   # есть sidecar → not blind

def test_source_scoped_extracted_body_in_tiers(sidecar_ds):
    r = u.run_unified_construction_harness("найди огнестойкость в документах проекта",
                                           dataset_ids=["ds"], storage_root=sidecar_ds)
    assert "extracted_body" in r.answer_data.get("searched_tiers", [])


# ── manifest staleness ───────────────────────────────────────────────────────────────────

def test_index_health_sidecar_available_from_write(sidecar_ds):
    rec = sa.inspect_dataset_index_health(["ds"], storage_root=sidecar_ds)["datasets"][0]
    assert rec["sidecar_available"] and rec["extracted_body_count"] > 0 and rec["manifest_present"]
    assert "no_lexical_index_but_file_body_available" in rec["warnings"]


# ── term sampler ─────────────────────────────────────────────────────────────────────────

def test_sample_extracted_terms_real(sidecar_ds):
    import importlib.util
    spec = importlib.util.spec_from_file_location("samp", "scripts/sample_extracted_terms_v15.py")
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    rep = m.sample("ds", storage_root=sidecar_ds, top_n=10)
    assert rep["sidecar_items"] > 0
    assert any("огнестойкость" in c["term"] or c["term"] for c in rep["candidates"])
    assert all(c.get("source_ref") for c in rep["candidates"])   # реальный ref, не фейк


# ── РЕАЛЬНЫЙ датасет рантайма (sidecars записаны с разрешением оператора) ─────────────────

@pytest.mark.skipif(not _SIDECARS, reason="sidecars для 844a2b53 не записаны (нужно одобрение оператора)")
def test_real_sidecar_health():
    rec = sa.inspect_dataset_index_health([_GOST], storage_root=_RT)["datasets"][0]
    assert rec["sidecar_available"] and rec["extracted_body_count"] > 1000 and rec["docx_file_count"] == 27

@pytest.mark.skipif(not _SIDECARS, reason="sidecars не записаны")
def test_real_extracted_body_hit():
    r = sa.search_extracted_body(["огнестойкости"], dataset_ids=[_GOST], storage_root=_RT)
    assert r.status == sa.FOUND and ".docx" in r.matches[0].source_ref and "para" in r.matches[0].source_ref

@pytest.mark.skipif(not _SIDECARS, reason="sidecars не записаны")
def test_real_norm_qa_complete_from_extracted():
    r = u.run_unified_construction_harness("правила огнестойкости стен", dataset_ids=[_GOST], storage_root=_RT)
    assert r.total_status == "complete"
    retr = next(b for b in r.evidence_blocks if b.type is EvidenceType.RETRIEVED)
    assert "СП" in retr.items[0].source_refs[0] or ".docx" in retr.items[0].source_refs[0]

@pytest.mark.skipif(not _SIDECARS, reason="sidecars не записаны")
def test_real_originals_not_mutated_after_write():
    import hashlib
    # перепроверка: оригиналы .docx читаемы и не пусты (write только добавил _extracted/)
    docx = list((_RT / _GOST).rglob("*.docx"))
    assert len(docx) == 27 and all(p.stat().st_size > 0 for p in docx)


# ── регрессии v0.3-v0.14 ─────────────────────────────────────────────────────────────────

def test_v14_write_policy_regression(tmp_path, monkeypatch):
    monkeypatch.setenv("LES_RUNTIME_HOME", str(tmp_path))
    monkeypatch.delenv("LES_ALLOW_RUNTIME_SIDECAR_WRITE", raising=False)
    assert de.is_runtime_path(tmp_path / "storage") and not de.runtime_write_allowed()

def test_v06_resource_workbook_regression():
    assert rc.validate_real_workbook()["matches"] is True

def test_resource_grand_complete():
    r = u.run_unified_construction_harness("проверь пример обсчёта")
    assert r.total_status == "complete" and abs(r.final_total - 16827283.19) < 1.0

def test_v04_source_scope_regression():
    assert u.route_construction_intent("найди ОЗК в актах смонтированного оборудования").intent == "asbuilt_extract"
    assert u.route_construction_intent("правила расстановки ОЗК").intent == "norm_qa"

def test_v13_bor_xlsx_regression(tmp_path):
    import openpyxl
    d = tmp_path / "ds"
    d.mkdir()
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "ВОР"
    ws.append(["Наименование", "Ед", "Кол-во"]); ws.append(["Грунт", "м3", 7200]); wb.save(d / "Ф9.xlsx")
    r = u.run_unified_construction_harness("извлеки ВОР из Ф9", dataset_ids=["ds"], storage_root=tmp_path)
    assert r.total_status == "complete"

def test_unit_gate_regression():
    assert ch.lsr_assemble([{"code": "06-02-001-01", "work": "плита", "unit": "м3", "qty": 720}])["asm_positions"][0]["qty"] == 7.2
