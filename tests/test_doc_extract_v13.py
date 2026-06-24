"""Unified Construction Harness v0.13 — PDF/DOCX/XLSX body extraction (sidecar, read-only, no OCR).

Извлечённый текст → searchable с source_ref до original-файла/страницы/абзаца/строки. Нет sidecar →
actionable. PDF без текст-слоя → no_text_layer (не фейк). Оригиналы не меняются. Контракт сохранён.
"""

import importlib.util
from pathlib import Path

import pytest

from proxy.services import doc_extract_service as de
from proxy.services import source_adapters as sa
from proxy.services import unified_construction_harness_service as u
from proxy.services import construction_harness_service as ch
from proxy.services import resource_cost_service as rc
from proxy.services.evidence_contract import EvidenceType

MONEY = 1.0


def _pdf(path, text="Akt smontirovannogo OZK-1 ustanovlen"):
    import fitz
    doc = fitz.open()
    pg = doc.new_page()
    if text:
        pg.insert_text((72, 72), text)
    doc.save(str(path))
    doc.close()


def _docx(path, heading="Спецификация оборудования", para="Клапан ОЗК-1, 6 шт", table=None):
    from docx import Document
    d = Document()
    d.add_heading(heading, 1)
    d.add_paragraph(para)
    if table:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for ri, row in enumerate(table):
            for ci, c in enumerate(row):
                t.rows[ri].cells[ci].text = str(c)
    d.save(str(path))


def _xlsx(path, sheet="ВОР", rows=(("Наименование", "Ед", "Кол-во"), ("Разработка грунта", "м3", 7200),
                                   ("Гидроизоляция", "м2", 1500))):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet
    for r in rows:
        ws.append(list(r))
    wb.save(str(path))


@pytest.fixture
def ds_extracted(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _pdf(d / "Акт.pdf")
    _docx(d / "Спецификация.docx")
    _xlsx(d / "Ф9_ВОР.xlsx")
    for f in d.iterdir():
        r = de.extract_file(f, ds="ds", rel=f.name)
        if r.items:
            de.write_sidecar(tmp_path, "ds", f.name, r.items)
    return tmp_path


# ── extraction ───────────────────────────────────────────────────────────────────────────

def test_extract_pdf_text(tmp_path):
    _pdf(tmp_path / "a.pdf")
    r = de.extract_pdf_text(tmp_path / "a.pdf", ds="ds", rel="a.pdf")
    assert r.status == "ok" and r.items[0].page == 1 and "#p1" in r.items[0].source_ref

def test_extract_pdf_no_text_layer(tmp_path):
    _pdf(tmp_path / "scan.pdf", text="")        # пустая страница = нет текст-слоя
    r = de.extract_pdf_text(tmp_path / "scan.pdf", ds="ds", rel="scan.pdf")
    assert r.status == "no_text_layer" and not r.items

def test_extract_docx_paragraphs_and_table(tmp_path):
    _docx(tmp_path / "d.docx", table=[["Наименование", "Ед", "Кол-во"], ["Грунт", "м3", "7200"]])
    r = de.extract_docx(tmp_path / "d.docx", ds="ds", rel="d.docx")
    assert r.status == "ok"
    kinds = {i.source_kind for i in r.items}
    assert "docx_text" in kinds and "docx_table" in kinds

def test_extract_xlsx_rows(tmp_path):
    _xlsx(tmp_path / "f.xlsx")
    r = de.extract_xlsx_generic(tmp_path / "f.xlsx", ds="ds", rel="f.xlsx")
    assert r.status == "ok" and any("#ВОР!R" in i.source_ref for i in r.items)

def test_every_sidecar_item_has_source_ref(tmp_path):
    _docx(tmp_path / "d.docx")
    r = de.extract_docx(tmp_path / "d.docx", ds="ds", rel="d.docx")
    assert all(i.source_ref for i in r.items)

def test_sidecar_write_read_roundtrip(tmp_path):
    _xlsx(tmp_path / "ds" / "f.xlsx") if (tmp_path / "ds").mkdir() or True else None
    r = de.extract_xlsx_generic(tmp_path / "ds" / "f.xlsx", ds="ds", rel="f.xlsx")
    de.write_sidecar(tmp_path, "ds", "f.xlsx", r.items)
    back = de.read_sidecars(tmp_path, "ds")
    assert len(back) == len(r.items) and back[0]["source_ref"]


# ── sidecar search ───────────────────────────────────────────────────────────────────────

def test_search_extracted_pdf_term(ds_extracted):
    r = sa.search_extracted_body(["ОЗК"], dataset_ids=["ds"], storage_root=ds_extracted)
    assert r.status == sa.FOUND and r.matches[0].source_kind == sa.KIND_EXTRACTED

def test_extracted_source_ref_points_to_original(ds_extracted):
    r = sa.search_extracted_body(["ОЗК"], dataset_ids=["ds"], storage_root=ds_extracted)
    # ref содержит оригинальный файл, НЕ sidecar-путь
    assert all(".jsonl" not in m.source_ref and "_extracted" not in m.source_ref for m in r.matches)

def test_no_sidecar_returns_no_source(tmp_path):
    (tmp_path / "ds").mkdir()
    r = sa.search_extracted_body(["ОЗК"], dataset_ids=["ds"], storage_root=tmp_path)
    assert r.status == sa.NO_SOURCE and any("no_extracted_body" in w for w in r.warnings)


# ── source-scoped via extracted_body ─────────────────────────────────────────────────────

def test_source_scoped_extracted_body_tier(ds_extracted):
    r = u.run_unified_construction_harness("найди ОЗК в спецификации", dataset_ids=["ds"], storage_root=ds_extracted)
    assert r.total_status == "complete"
    assert "extracted_body" in r.answer_data.get("searched_tiers", []) or r.sources

def test_source_scoped_not_found_searches_extracted_tier(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _docx(d / "Спецификация.docx", para="Насос К-100")     # ОЗК нет
    r2 = de.extract_docx(d / "Спецификация.docx", ds="ds", rel="Спецификация.docx")
    de.write_sidecar(tmp_path, "ds", "Спецификация.docx", r2.items)
    r = u.run_unified_construction_harness("найди ОЗК в спецификации", dataset_ids=["ds"], storage_root=tmp_path)
    assert "extracted_body" in r.answer_data.get("searched_tiers", [])


# ── norm QA via extracted ────────────────────────────────────────────────────────────────

def test_norm_qa_from_extracted_docx(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _docx(d / "СП_котельная.docx", heading="СП котельная", para="АУПТ требуется для котельной по п.5.4")
    r = de.extract_docx(d / "СП_котельная.docx", ds="ds", rel="СП_котельная.docx")
    de.write_sidecar(tmp_path, "ds", "СП_котельная.docx", r.items)
    res = u.run_unified_construction_harness("найди АУПТ в документах проекта", dataset_ids=["ds"], storage_root=tmp_path)
    assert res.total_status == "complete"

def test_norm_qa_no_extracted_body_tiers(tmp_path):
    (tmp_path / "ds").mkdir()
    r = u.run_unified_construction_harness("правила расстановки ВВГнг", dataset_ids=["ds"], storage_root=tmp_path)
    assert r.total_status == "no_data"
    assert "extracted_body" in r.answer_data.get("searched_tiers", [])


# ── BOR/LSR from xlsx/docx ───────────────────────────────────────────────────────────────

def test_bor_from_xlsx(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _xlsx(d / "Ф9.xlsx")
    r = u.run_unified_construction_harness("извлеки ВОР из Ф9", dataset_ids=["ds"], storage_root=tmp_path)
    assert r.total_status == "complete"
    assert next(b for b in r.evidence_blocks if b.type is EvidenceType.RETRIEVED).items[0].source_refs

def test_bor_from_docx_table(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _docx(d / "ВОР.docx", table=[["Наименование", "Ед", "Кол-во"], ["Разработка грунта", "м3", "7200"],
                                 ["Гидроизоляция", "м2", "1500"]])
    r = u.run_unified_construction_harness("извлеки ВОР из Ф9", dataset_ids=["ds"], storage_root=tmp_path)
    assert r.total_status == "complete"

def test_lsr_from_xlsx_bor(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _xlsx(d / "Ф9.xlsx", rows=(("Наименование", "Ед", "Кол-во"),
                               ("разработка грунта в котловане", "м3", 7200),
                               ("устройство монолитной фундаментной плиты", "м3", 720)))
    r = u.run_unified_construction_harness("собери предварительную ЛСР по Ф9", dataset_ids=["ds"], storage_root=tmp_path)
    assert r.total_status in ("complete", "partial")        # ЛСР пошла из xlsx-строк


# ── index health v0.13 ───────────────────────────────────────────────────────────────────

def test_index_health_binary_counts(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _pdf(d / "a.pdf")
    _docx(d / "b.docx")
    _xlsx(d / "c.xlsx")
    rec = sa.inspect_dataset_index_health(["ds"], storage_root=tmp_path)["datasets"][0]
    assert rec["pdf_file_count"] == 1 and rec["docx_file_count"] == 1 and rec["xlsx_file_count"] == 1

def test_index_health_pdf_without_sidecar_warning(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _pdf(d / "a.pdf")          # есть PDF, sidecar нет
    rec = sa.inspect_dataset_index_health(["ds"], storage_root=tmp_path)["datasets"][0]
    assert "pdf_files_without_sidecars" in rec["warnings"] and "no_extracted_body" in rec["warnings"]

def test_index_health_sidecar_available(ds_extracted):
    rec = sa.inspect_dataset_index_health(["ds"], storage_root=ds_extracted)["datasets"][0]
    assert rec["sidecar_available"] and rec["extracted_body_count"] > 0
    assert "no_extracted_body" not in rec["warnings"]


# ── extraction script (dry-run / path-safe) ──────────────────────────────────────────────

def _load_script():
    spec = importlib.util.spec_from_file_location("extract_v13", "scripts/extract_dataset_bodies_v13.py")
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    return m

def test_extract_script_dry_run_no_write(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _xlsx(d / "f.xlsx")
    s = _load_script()
    rep = s.run("ds", storage_root=tmp_path, exts={".xlsx"}, max_files=10, max_mb=40, dry_run=True, force=False)
    assert rep["files_extracted"] == 1 and rep["xlsx_rows"] > 0
    assert de.sidecar_count(tmp_path, "ds") == 0          # dry-run НЕ пишет

def test_extract_script_writes(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _xlsx(d / "f.xlsx")
    s = _load_script()
    s.run("ds", storage_root=tmp_path, exts={".xlsx"}, max_files=10, max_mb=40, dry_run=False, force=False)
    assert de.sidecar_count(tmp_path, "ds") == 1

def test_extract_script_path_traversal_safe(tmp_path):
    s = _load_script()
    assert s._safe_under(tmp_path / "ds", tmp_path / "ds" / "f.xlsx") is True
    assert s._safe_under(tmp_path / "ds", tmp_path / "other" / "f.xlsx") is False


# ── safety: оригиналы не меняются ────────────────────────────────────────────────────────

def test_originals_not_mutated(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    _xlsx(d / "f.xlsx")
    before = (d / "f.xlsx").read_bytes()
    de.extract_xlsx_generic(d / "f.xlsx", ds="ds", rel="f.xlsx")
    s = _load_script()
    s.run("ds", storage_root=tmp_path, exts={".xlsx"}, max_files=10, max_mb=40, dry_run=False, force=False)
    assert (d / "f.xlsx").read_bytes() == before


# ── регрессии v0.3-v0.12 ─────────────────────────────────────────────────────────────────

def test_v12_file_body_md_regression(tmp_path):
    d = tmp_path / "ds"
    d.mkdir()
    (d / "x.md").write_text("Клапан ОЗК-1 установлен\n", encoding="utf-8")
    assert sa.search_file_body(["ОЗК"], dataset_ids=["ds"], storage_root=tmp_path).status == sa.FOUND

def test_v06_resource_workbook_regression():
    assert rc.validate_real_workbook()["matches"] is True

def test_resource_grand_complete():
    r = u.run_unified_construction_harness("проверь пример обсчёта")
    assert r.total_status == "complete" and abs(r.final_total - 16827283.19) < MONEY

def test_v04_source_scope_regression():
    assert u.route_construction_intent("найди ОЗК в актах смонтированного оборудования").intent == "asbuilt_extract"
    assert u.route_construction_intent("правила расстановки ОЗК").intent == "norm_qa"

def test_v03_lsr_parquet_regression(tmp_path):
    dsid = ch.write_demo_project_doc(tmp_path)
    r = u.run_unified_construction_harness("собери предварительную ЛСР по Ф9", dataset_ids=[dsid], storage_root=tmp_path)
    assert r.total_status == "complete" and r.final_total is not None
