from __future__ import annotations

import zipfile

import pytest

from proxy.services.file_viewer_service import (
    FILE_VIEWER_SCHEMA,
    file_viewer_html,
    is_viewable_file,
    viewer_kind,
)


def test_text_viewer_escapes_active_markup(tmp_path):
    source = tmp_path / "unsafe.html"
    source.write_text('<script>alert("x")</script><b>text</b>', encoding="utf-8")

    output = file_viewer_html(source, path_id="docs/unsafe.html")

    assert FILE_VIEWER_SCHEMA in output
    assert "&lt;script&gt;" in output
    assert '<script>alert("x")</script>' not in output
    assert str(tmp_path) not in output


def test_pptx_viewer_reads_slide_text_without_office(tmp_path):
    source = tmp_path / "meeting.pptx"
    slide = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
      xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>Decision 42</a:t></a:r></a:p>
      </p:txBody></p:sp></p:spTree></p:cSld></p:sld>'''
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr("ppt/slides/slide1.xml", slide)

    output = file_viewer_html(source, path_id="docs/meeting.pptx", locator="slide1")

    assert "Decision 42" in output
    assert "Слайд 1" in output
    assert "evidence-hit" in output


def test_eml_viewer_uses_plain_body_and_escapes_html(tmp_path):
    source = tmp_path / "message.eml"
    source.write_bytes(
        "From: engineer@example.test\nTo: pm@example.test\nSubject: Check\n"
        "Content-Type: text/plain; charset=utf-8\n\n<confirm>Да</confirm>".encode()
    )

    output = file_viewer_html(source, path_id="mail/message.eml")

    assert "engineer@example.test" in output
    assert "&lt;confirm&gt;" in output
    assert "<confirm>" not in output


def test_docx_viewer_preserves_paragraph_table_order(tmp_path):
    from docx import Document

    source = tmp_path / "ordered.docx"
    document = Document()
    document.add_paragraph("До таблицы")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "В таблице"
    document.add_paragraph("После таблицы")
    document.save(source)

    output = file_viewer_html(source, path_id="docs/ordered.docx")

    assert output.index("До таблицы") < output.index("В таблице") < output.index("После таблицы")


@pytest.mark.parametrize(
    ("name", "kind"),
    [("a.pdf", "pdf"), ("a.docx", "word"), ("a.xlsx", "excel"),
     ("a.pptx", "presentation"), ("a.png", "image"), ("a.md", "text")],
)
def test_viewer_kind_contract(name, kind):
    assert is_viewable_file(name)
    assert viewer_kind(name) == kind


def test_legacy_office_is_not_misrepresented_as_modern_format():
    assert not is_viewable_file("legacy.doc")
    assert not is_viewable_file("legacy.xls")
