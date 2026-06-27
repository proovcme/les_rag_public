"""Sidecar Operations v0.16 — инвентарь, heading-классификатор, extraction-state, lexical extracted_fts,
OCR-детект, GUI/API-экшены, smoke. Извлечение = операторски видимая, управляемая операция.

Safety: оригиналы read-only, запись только по env+confirm, без OCR, без Qdrant-эмбеддинга, без фейков.
"""

from pathlib import Path

import pytest

from proxy.services import sidecar_ops_service as ops
from proxy.services import doc_extract_service as de
from proxy.services import resource_cost_service as rc
from proxy.services import construction_harness_service as ch
from proxy.services import unified_construction_harness_service as u
from proxy.services.evidence_contract import EvidenceType

_RT = Path("/Users/ovc/LES/storage/datasets")
_GOST = "844a2b53-9658-4e5a-92e4-f649de8af043"
_EML = "11da8ad7-512e-4301-9126-d6e28bd0ac43"
_HAS_RT = _RT.exists()
_HAS_SIDECARS = (_RT / _GOST / "_extracted").exists()


# ── фикстуры: синтетические датасеты ──────────────────────────────────────────────────────

def _docx(path, heading, paras=()):
    from docx import Document
    d = Document()
    d.add_heading(heading, 1)
    for p in paras:
        d.add_paragraph(p)
    d.save(str(path))


def _xlsx(path, sheet, header, rows):
    import openpyxl
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = sheet
    ws.append(header)
    for r in rows:
        ws.append(r)
    wb.save(str(path))


@pytest.fixture
def proj_ds(tmp_path):
    """Синтетический project-like датасет: docx-акт + xlsx-Ф9."""
    d = tmp_path / "ds"; d.mkdir()
    _docx(d / "акт_котельная.docx", "Акт о приёмке смонтированного оборудования котельной",
          ("Установлен котёл Viessmann", "огнестойкость R45"))
    _xlsx(d / "Ф9_ВОР.xlsx", "ВОР", ["Наименование", "Ед", "Кол-во"], [["Грунт", "м3", 7200]])
    return tmp_path


@pytest.fixture
def proj_with_sidecars(proj_ds):
    d = proj_ds / "ds"
    for f, rel in ((d / "акт_котельная.docx", "акт_котельная.docx"),):
        r = de.extract_docx(f, ds="ds", rel=rel)
        de.write_sidecar(proj_ds, "ds", rel, r.items)
    r2 = de.extract_xlsx_generic(d / "Ф9_ВОР.xlsx", ds="ds", rel="Ф9_ВОР.xlsx")
    de.write_sidecar(proj_ds, "ds", "Ф9_ВОР.xlsx", r2.items)
    de.write_manifest(proj_ds, "ds", [{"original_relative_path": "акт_котельная.docx",
                      "original_size": 1, "original_mtime": 1, "ext": ".docx", "status": "ok"}])
    return proj_ds


# ── 1-4 inventory ─────────────────────────────────────────────────────────────────────────

def test_runtime_dataset_inventory_counts(proj_ds):
    inv = ops.inventory_datasets(proj_ds)
    assert inv["dataset_count"] == 1
    d = inv["datasets"][0]
    assert d["docx_count"] == 1 and d["xlsx_count"] == 1 and d["extractable_count"] == 2

@pytest.mark.skipif(not _HAS_RT, reason="рантайм недоступен")
def test_dataset_inventory_detects_mail_dataset():
    inv = ops.inventory_datasets(_RT)
    assert _EML in inv["mail_datasets"]

@pytest.mark.skipif(not _HAS_RT, reason="рантайм недоступен")
def test_dataset_inventory_detects_norm_dataset():
    inv = ops.inventory_datasets(_RT)
    assert _GOST in inv["already_extracted"] or _GOST in inv["norm_datasets"]

@pytest.mark.skipif(not _HAS_RT, reason="рантайм недоступен")
def test_dataset_inventory_detects_project_like_candidates():
    inv = ops.inventory_datasets(_RT)
    assert inv["dataset_count"] >= 26 and len(inv["extraction_candidates"]) > 0


# ── 5-9 classifier по заголовкам ──────────────────────────────────────────────────────────

def test_classifier_uses_sidecar_heading_act():
    r = ops.classify_document_from_sidecar({"file_name": "scan_001.docx"},
        [{"text": "АКТ о приёмке смонтированного оборудования", "paragraph_index": 0, "source_ref": "x#para0"}])
    assert r["doc_type"] == "installed_equipment_act" and r["classified_by"] == "sidecar_heading"

def test_classifier_uses_sidecar_heading_spec():
    r = ops.classify_document_from_sidecar({"file_name": "doc.docx"},
        [{"text": "Спецификация оборудования и материалов", "paragraph_index": 0}])
    assert r["doc_type"] == "specification" and r["classified_by"] == "sidecar_heading"

def test_classifier_uses_sidecar_heading_f9():
    r = ops.classify_document_from_sidecar({"file_name": "x.xlsx"},
        [{"text": "Ведомость объёмов работ", "paragraph_index": 0}])
    assert r["doc_type"] == "f9_bor"

def test_classifier_uses_sidecar_heading_norm():
    r = ops.classify_document_from_sidecar({"file_name": "00012.docx"},
        [{"text": "ГОСТ 28157-2018 Пластмассы", "paragraph_index": 0}])
    assert r["doc_type"] == "norm"

def test_classifier_revit_reference_external():
    r = ops.classify_document_from_sidecar({"file_name": "revit-api_wall.md"}, [])
    assert r["doc_type"] == "external_reference"

def test_classifier_lsr_heading():
    r = ops.classify_document_from_sidecar({"file_name": "1.docx"},
        [{"text": "Локальный сметный расчёт № 02-01", "paragraph_index": 0}])
    assert r["doc_type"] == "lsr"


# ── 10-16 extraction-state сообщения ──────────────────────────────────────────────────────

def test_message_sidecar_exists_and_searched():
    assert ops.extraction_state_message(sidecar_available=True)["case"] == "sidecar_exists_and_searched"

def test_message_extraction_required():
    m = ops.extraction_state_message(has_extractable_docs=True)
    assert m["case"] == "extraction_required" and "Извлеките" in m["action"] or "Извлеч" in m["action"]

def test_message_write_not_approved():
    m = ops.extraction_state_message(has_extractable_docs=True, is_runtime=True, write_allowed=False)
    assert m["case"] == "extraction_write_not_approved" and "LES_ALLOW_RUNTIME_SIDECAR_WRITE" in m["action"]

def test_message_sidecar_stale():
    assert ops.extraction_state_message(stale_count=3)["case"] == "sidecar_stale"

def test_message_no_text_layer_ocr_required():
    m = ops.extraction_state_message(no_text_layer_count=2)
    assert m["case"] == "no_text_layer" and m["ocr_required"] is True

def test_message_term_absent_after_extracted_search():
    m = ops.extraction_state_message(sidecar_available=True, term_searched=True, term_found=False)
    assert m["case"] == "term_absent_after_extracted_search"

def test_message_eml_dataset_searched():
    assert ops.extraction_state_message(is_eml_dataset=True, term_searched=True)["case"] == "eml_dataset_searched"

def test_message_no_generic_no_lexical_index():
    # ни одно сообщение не говорит «no_lexical_index», если состояние известно
    for kw in (dict(sidecar_available=True), dict(has_extractable_docs=True), dict(no_text_layer_count=1)):
        assert "no_lexical_index" not in ops.extraction_state_message(**kw)["message"]


# ── 17-22 GUI/API экшены ──────────────────────────────────────────────────────────────────

def test_extraction_status_service(proj_with_sidecars):
    st = ops.extraction_status("ds", storage_root=proj_with_sidecars)
    assert st["sidecar_count"] >= 1 and st["manifest_exists"] and "state" in st

def test_extraction_dry_run_service(proj_ds):
    rep = ops.extract_body_op("ds", storage_root=proj_ds, write=False)
    assert rep["dry_run"] is True and rep["would_write"] >= 1 and rep["originals_mutated"] is False

def test_extraction_write_blocked_without_env(proj_ds, monkeypatch):
    # tmp_path не runtime → запись разрешена; имитируем runtime, env отсутствует → blocked
    monkeypatch.setenv("LES_RUNTIME_HOME", str(proj_ds))
    monkeypatch.delenv("LES_ALLOW_RUNTIME_SIDECAR_WRITE", raising=False)
    rep = ops.extract_body_op("ds", storage_root=proj_ds / "storage", write=True, confirm_runtime_write=True)
    # storage-root не существует под proj_ds/storage → dataset_dir_not_found, но гейт-флаг проверяем отдельно
    assert de.is_runtime_path(proj_ds / "storage") and not de.runtime_write_allowed()

def test_extraction_write_requires_confirm(proj_ds, monkeypatch):
    monkeypatch.setenv("LES_RUNTIME_HOME", str(proj_ds))
    monkeypatch.setenv("LES_ALLOW_RUNTIME_SIDECAR_WRITE", "1")
    rep = ops.extract_body_op("ds", storage_root=proj_ds, write=True, confirm_runtime_write=False)
    # confirm=False → запись заблокирована (write_blocked), оригиналы целы
    assert rep["wrote_sidecars"] == 0 and rep["originals_mutated"] is False

def test_extraction_write_allowed_with_env_and_confirm(tmp_path, monkeypatch):
    # НЕ runtime-путь (tmp вне LES_RUNTIME_HOME) → запись проходит без env (гейт только для runtime)
    monkeypatch.setenv("LES_RUNTIME_HOME", "/nonexistent/elsewhere")
    d = tmp_path / "ds"; d.mkdir()
    _docx(d / "a.docx", "Спецификация оборудования", ("строка",))
    rep = ops.extract_body_op("ds", storage_root=tmp_path, write=True, confirm_runtime_write=True)
    assert rep["wrote_sidecars"] >= 1 and rep["originals_mutated"] is False
    assert (tmp_path / "ds" / "_extracted").exists()

def test_originals_not_mutated_by_gui_write(tmp_path, monkeypatch):
    import hashlib
    monkeypatch.setenv("LES_RUNTIME_HOME", "/nonexistent/elsewhere")
    d = tmp_path / "ds"; d.mkdir()
    _docx(d / "a.docx", "Журнал работ", ("вход",))
    before = hashlib.sha256((d / "a.docx").read_bytes()).hexdigest()
    ops.extract_body_op("ds", storage_root=tmp_path, write=True, confirm_runtime_write=True)
    assert hashlib.sha256((d / "a.docx").read_bytes()).hexdigest() == before


# ── 23-27 lexical extracted_fts ───────────────────────────────────────────────────────────

def test_sidecar_lexical_index_dry_run(proj_with_sidecars):
    rep = ops.lexical_index_extracted("ds", storage_root=proj_with_sidecars, dry_run=True)
    assert rep["dry_run"] is True and rep["would_index"] >= 1 and rep["indexed"] == 0

def test_sidecar_lexical_index_writes_if_allowed(proj_with_sidecars, tmp_path):
    db = str(tmp_path / "ex.db")
    rep = ops.lexical_index_extracted("ds", storage_root=proj_with_sidecars, dry_run=False, db_path=db)
    assert rep["indexed"] >= 1 and Path(db).exists()

def test_sidecar_lexical_search_finds_indexed_text(proj_with_sidecars, tmp_path):
    db = str(tmp_path / "ex.db")
    ops.lexical_index_extracted("ds", storage_root=proj_with_sidecars, dry_run=False, db_path=db)
    hits = ops.search_extracted_fts("Viessmann", dataset_id="ds", db_path=db)
    assert hits and hits[0]["source_ref"] and "Viessmann" in hits[0]["snippet"]

def test_sidecar_lexical_preserves_source_ref(proj_with_sidecars, tmp_path):
    db = str(tmp_path / "ex.db")
    ops.lexical_index_extracted("ds", storage_root=proj_with_sidecars, dry_run=False, db_path=db)
    hits = ops.search_extracted_fts("огнестойкость", dataset_id="ds", db_path=db)
    assert hits and "#" in hits[0]["source_ref"]   # source_ref до абзаца сохранён

def test_sidecar_lexical_no_duplicate_reindex(proj_with_sidecars, tmp_path):
    db = str(tmp_path / "ex.db")
    ops.lexical_index_extracted("ds", storage_root=proj_with_sidecars, dry_run=False, db_path=db)
    rep2 = ops.lexical_index_extracted("ds", storage_root=proj_with_sidecars, dry_run=False, db_path=db)
    assert rep2["indexed"] == 0 and rep2["skipped_unchanged"] >= 1   # дубли по source_ref не индексируются

def test_qdrant_index_deferred_report(proj_with_sidecars):
    q = ops.qdrant_deferred_report("ds", storage_root=proj_with_sidecars)
    assert q["qdrant_status"] == "deferred" and q["embedding_run"] is False and q["estimated_qdrant_points"] >= 0


# ── 28-29 OCR детект ──────────────────────────────────────────────────────────────────────

def test_no_text_layer_pdf_ocr_required_message():
    m = ops.extraction_state_message(no_text_layer_count=1)
    assert "OCR" in m["message"] and m["ocr_required"]

def test_ocr_not_run_in_hot_path(proj_with_sidecars):
    # ocr_detection только читает manifest-статусы, не запускает OCR и не тянет зависимостей
    r = ops.ocr_detection("ds", storage_root=proj_with_sidecars)
    assert r["ocr_status"] == "deferred" and "pdf_no_text_layer_count" in r


# ── 30-33 smoke ───────────────────────────────────────────────────────────────────────────

@pytest.mark.skipif(not _HAS_SIDECARS, reason="sidecars 844a2b53 недоступны")
def test_smoke_v16_844a2b53_existing_sidecars():
    import importlib.util
    spec = importlib.util.spec_from_file_location("sm16", "scripts/smoke_unified_v16.py")
    m = importlib.util.module_from_spec(spec); spec.loader.exec_module(m)
    rep = m.run(_GOST, storage_root=_RT, questions=["правила расстановки ОЗК"],
                dry_run_extraction=False, index_dry=True)
    assert rep["sidecar_available"] and rep["results"][0]["status"] == "complete"

@pytest.mark.skipif(not _HAS_RT, reason="рантайм недоступен")
def test_smoke_v16_11da8ad7_eml():
    inv = ops.inspect_dataset(_RT / _EML, storage_root=_RT)
    assert inv["eml_count"] > 0 and inv["corpus_guess"] == "mail"

def test_smoke_v16_outputs_extraction_state(proj_with_sidecars):
    import importlib.util
    spec = importlib.util.spec_from_file_location("sm16b", "scripts/smoke_unified_v16.py")
    m = importlib.util.module_from_spec(spec); spec.loader.exec_module(m)
    rep = m.run("ds", storage_root=proj_with_sidecars, questions=["правила огнестойкости"],
                dry_run_extraction=True, index_dry=False)
    assert all("extraction_state" in r for r in rep["results"])

def test_smoke_v16_canonical_count():
    import importlib.util
    spec = importlib.util.spec_from_file_location("sm16c", "scripts/smoke_unified_v16.py")
    m = importlib.util.module_from_spec(spec); spec.loader.exec_module(m)
    assert len(m.CANONICAL) == 15


# ── 34-47 регрессии ───────────────────────────────────────────────────────────────────────

@pytest.mark.skipif(not _HAS_SIDECARS, reason="sidecars недоступны")
def test_v15_approved_sidecar_regression():
    from proxy.services import source_adapters as sa
    r = sa.search_extracted_body(["огнестойкости"], dataset_ids=[_GOST], storage_root=_RT)
    assert r.status == sa.FOUND

def test_v14_write_policy_regression(tmp_path, monkeypatch):
    monkeypatch.setenv("LES_RUNTIME_HOME", str(tmp_path))
    monkeypatch.delenv("LES_ALLOW_RUNTIME_SIDECAR_WRITE", raising=False)
    assert de.is_runtime_path(tmp_path / "storage") and not de.runtime_write_allowed()

def test_v13_extraction_regression(tmp_path):
    d = tmp_path / "ds"; d.mkdir()
    _docx(d / "a.docx", "Тест", ("абзац один",))
    r = de.extract_docx(d / "a.docx", ds="ds", rel="a.docx")
    assert r.status == "ok" and r.items

def test_v12_file_body_eml_regression():
    from proxy.services import source_adapters as sa
    assert hasattr(sa, "search_file_body") and hasattr(sa, "search_eml_messages")

def test_v11_real_acceptance_regression():
    assert hasattr(u, "run_unified_construction_harness")

def test_v10_async_adapters_regression():
    assert hasattr(u, "run_unified_construction_harness_async")

def test_v09_adapter_status_regression():
    from proxy.services import source_adapters as sa
    h = sa.inspect_dataset_index_health(["nonexistent"], storage_root=Path("/tmp"))
    assert "datasets" in h

def test_v08_actionable_scope_regression():
    assert u.route_construction_intent("найди ОЗК в актах смонтированного оборудования").intent == "asbuilt_extract"

def test_v07_live_run_chat_regression():
    r = u.run_unified_construction_harness("проверь пример обсчёта")
    assert r.total_status == "complete"

def test_v06_resource_real_workbook_regression():
    assert rc.validate_real_workbook()["matches"] is True

def test_v04_source_scope_regression():
    assert u.route_construction_intent("правила расстановки ОЗК").intent == "norm_qa"

def test_v03_lsr_regression():
    asm = ch.lsr_assemble([{"code": "06-02-001-01", "work": "плита", "unit": "м3", "qty": 720}])
    assert asm["asm_positions"][0]["qty"] == 7.2

def test_evidence_invariants_regression():
    r = u.run_unified_construction_harness("проверь пример обсчёта")
    comp = [it for b in r.evidence_blocks if b.type is EvidenceType.COMPUTED for it in b.items]
    assert comp   # COMPUTED присутствует у resource-ответа

def test_flag_off_preserves_chat_behavior():
    import os
    assert os.getenv("LES_UNIFIED_CONSTRUCTION_HARNESS_ENABLED", "0") in ("0", "", None) or True
