"""W11.3/W19 — сервис типовых форм: дескриптор + данные объекта → документ (0 LLM)."""
import importlib

import pytest


@pytest.fixture()
def env(tmp_path, monkeypatch):
    (tmp_path / "data").mkdir()
    forms_dir = tmp_path / "forms"
    forms_dir.mkdir()
    (forms_dir / "templates").mkdir()
    # минимальный дескриптор формы для теста
    (forms_dir / "demo.yaml").write_text(
        "id: demo\n"
        "title: Демо-акт\n"
        "legal_basis: Тестовое основание\n"
        "templates: { docx: " + str(forms_dir / "templates" / "demo.docx") + " }\n"
        "fields:\n"
        "  - { key: object_name, label: Объект, type: str, source: \"project.name\" }\n"
        "  - { key: volume,      label: Объём, type: num, source: \"field.total\" }\n"
        "  - { key: ntd,         label: НТД,   type: list, source: \"edges.references_ntd\" }\n"
        "  - { key: note,        label: Прим., type: str, source: manual }\n"
        "  - { key: today,       label: Дата,  type: str, source: date.today }\n",
        encoding="utf-8",
    )
    monkeypatch.setenv("RAG_META_DB_PATH", str(tmp_path / "data" / "les_meta.db"))
    monkeypatch.setenv("LES_FORMS_DIR", str(forms_dir))
    monkeypatch.setenv("LES_FORMS_OUT_DIR", str(tmp_path / "data" / "forms_out"))
    import backend.rag_config as rc
    importlib.reload(rc)
    for mod in ("proxy.services.project_service", "proxy.services.field_intake_service",
                "proxy.services.edge_service", "proxy.services.forms_service"):
        importlib.reload(importlib.import_module(mod))
    import proxy.services.forms_service as fs
    return fs, forms_dir, tmp_path


def _seed_project(name="БЦ Тест"):
    from proxy.services.project_service import create_project
    return create_project(name, code="T", address="ул. Тестовая")["id"]


# ── реестр + резолв ──────────────────────────────────────────────────

def test_list_and_load(env):
    fs, *_ = env
    forms = fs.list_forms()
    assert any(f["id"] == "demo" for f in forms)
    assert fs.load_descriptor("missing") is None


def test_resolve_from_project_and_services(env):
    fs, *_ = env
    pid = _seed_project("БЦ «Банкрот»")
    from proxy.services.field_intake_service import create_entry
    create_entry("Бетон", 12.5, "м3", status="confirmed")
    create_entry("Бетон", 7.5, "м3", status="confirmed")
    from proxy.services.edge_service import add_edge
    add_edge("note", "1", "ntd_code", "СП 48.13330", "references_ntd", method="regex_ntd")

    resolved = fs.resolve_fields("demo", project_id=pid, manual={"note": "вручную"})
    vals = {f["key"]: f["value"] for f in resolved["fields"]}
    assert vals["object_name"] == "БЦ «Банкрот»"  # project.name
    assert vals["volume"] == "20"                  # field.total = 12.5+7.5, SQL не LLM
    assert vals["ntd"] == "СП 48.13330"            # edges.references_ntd
    assert vals["note"] == "вручную"               # manual override
    assert vals["today"]                           # date.today заполнен


def test_manual_needs_input_flag(env):
    fs, *_ = env
    pid = _seed_project()
    resolved = fs.resolve_fields("demo", project_id=pid, manual={})
    note = next(f for f in resolved["fields"] if f["key"] == "note")
    assert note["needs_input"] is True


def test_resolve_unknown_form(env):
    fs, *_ = env
    assert fs.resolve_fields("nope") is None


# ── рендереры ────────────────────────────────────────────────────────

def test_render_html(env):
    fs, *_ = env
    pid = _seed_project("Объект-Х")
    html = fs.generate("demo", "html", project_id=pid)["html"]
    assert "Демо-акт" in html and "Объект-Х" in html


def test_render_docx_template_less(env):
    fs, _, tmp_path = env
    pid = _seed_project("Объект-Д")
    res = fs.generate("demo", "docx", project_id=pid, manual={"note": "примечание-X"})
    from docx import Document
    doc = Document(res["path"])
    text = "\n".join(p.text for p in doc.paragraphs)
    cell_text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Демо-акт" in text
    assert "Объект-Д" in cell_text  # значение из проекта в таблице
    assert "примечание-X" in cell_text


def test_render_docx_fills_template(env):
    fs, forms_dir, _ = env
    # создаём образец с метками {{...}}
    from docx import Document
    tmpl = Document()
    tmpl.add_paragraph("Объект: {{object_name}}")
    table = tmpl.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Объём:"
    table.rows[0].cells[1].text = "{{volume}} м3"
    tmpl.save(str(forms_dir / "templates" / "demo.docx"))

    pid = _seed_project("Объект-Шаблон")
    from proxy.services.field_intake_service import create_entry
    create_entry("Бетон", 42.0, "м3", status="confirmed")
    res = fs.generate("demo", "docx", project_id=pid)

    doc = Document(res["path"])
    text = "\n".join(p.text for p in doc.paragraphs)
    cell_text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Объект: Объект-Шаблон" in text   # {{object_name}} подставлен
    assert "42 м3" in cell_text              # {{volume}} подставлен в ячейку
    assert "{{" not in text and "{{" not in cell_text  # плейсхолдеров не осталось


def test_render_xlsx(env):
    fs, *_ = env
    pid = _seed_project("Объект-E")
    res = fs.generate("demo", "xlsx", project_id=pid)
    from openpyxl import load_workbook
    wb = load_workbook(res["path"])
    ws = wb.active
    assert ws["A1"].value == "Демо-акт"
    flat = [c.value for row in ws.iter_rows() for c in row]
    assert "Объект-E" in flat


def test_generate_rejects_bad_fmt(env):
    fs, *_ = env
    pid = _seed_project()
    with pytest.raises(ValueError):
        fs.generate("demo", "pdf", project_id=pid)
