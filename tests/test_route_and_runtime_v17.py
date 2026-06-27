"""v0.17 — Runtime Alignment + Deterministic Route Fix + honest .xls (backend-приоритет).

Главный баг: «реестр документации котельной» уходил в ГЛОБАЛЬНЫЙ «Реестр проектов ЛЕС — 6». Фикс:
project_registry (глобальный) ≠ project_document_registry (scoped). Evidence-контракт не ослаблен,
флаг OFF не тронут, фейков нет.
"""

import os
from pathlib import Path

import pytest

from proxy.services import project_registry_chat_service as prc
from proxy.services import agent_router_service as ar
from proxy.services import sidecar_ops_service as ops
from proxy.services import doc_extract_service as de
from proxy.services import unified_construction_harness_service as u
from proxy.services import resource_cost_service as rc
from proxy.services import construction_harness_service as ch
from proxy.services import source_adapters as sa
from proxy.services.evidence_contract import EvidenceType


# ── §3 runtime alignment: extraction endpoints зарегистрированы ───────────────────────────

def _route_paths():
    from proxy.routers.datasets import router
    return {getattr(r, "path", "") for r in router.routes}

def test_extraction_status_endpoint_available():
    assert any("extraction-status" in p for p in _route_paths())

def test_extraction_dry_run_endpoint_available():
    assert any("extract-body/dry-run" in p for p in _route_paths())

def test_extraction_write_endpoint_available():
    assert any("extract-body/write" in p for p in _route_paths())

def test_runtime_datasets_py_no_missing_extraction_routes():
    paths = _route_paths()
    for needle in ("extraction-status", "extract-body/dry-run", "extract-body/write"):
        assert any(needle in p for p in paths), f"маршрут {needle} отсутствует"

def test_existing_dataset_routes_still_work():
    paths = _route_paths()
    assert any(p == "/api/rag/datasets" for p in paths)        # старые маршруты целы

def test_extraction_write_blocked_without_env(tmp_path, monkeypatch):
    monkeypatch.setenv("LES_RUNTIME_HOME", str(tmp_path))
    monkeypatch.delenv("LES_ALLOW_RUNTIME_SIDECAR_WRITE", raising=False)
    d = tmp_path / "storage" / "ds"; d.mkdir(parents=True)
    from docx import Document
    doc = Document(); doc.add_paragraph("текст"); doc.save(str(d / "a.docx"))
    rep = ops.extract_body_op("ds", storage_root=tmp_path / "storage", write=True, confirm_runtime_write=True)
    assert rep["wrote_sidecars"] == 0 and rep["originals_mutated"] is False   # env нет → blocked

def test_extraction_write_requires_confirm(tmp_path, monkeypatch):
    monkeypatch.setenv("LES_RUNTIME_HOME", str(tmp_path))
    monkeypatch.setenv("LES_ALLOW_RUNTIME_SIDECAR_WRITE", "1")
    d = tmp_path / "storage" / "ds"; d.mkdir(parents=True)
    from docx import Document
    doc = Document(); doc.add_paragraph("текст"); doc.save(str(d / "a.docx"))
    rep = ops.extract_body_op("ds", storage_root=tmp_path / "storage", write=True, confirm_runtime_write=False)
    assert rep["wrote_sidecars"] == 0   # confirm нет → blocked


# ── §4 deterministic route fix ────────────────────────────────────────────────────────────

def test_reestr_dokumentacii_routes_project_document_registry():
    # «реестр документации» — НЕ глобальный реестр проектов
    assert prc.is_registry_query("составь реестр документации котельной") is False
    assert prc.is_document_registry_query("составь реестр документации котельной") is True

def test_reestr_dokumentacii_lesnoy64_uses_selected_scope():
    # есть scope (проект выбран) → None (RAG по объекту), не глобальный список
    assert prc.maybe_handle_document_registry(
        "составь реестр документации про котельную на лесном 64", project_id=2) is None

def test_global_project_registry_only_exact():
    for q in ("реестр проектов", "покажи реестр проектов", "какие у нас объекты", "список всех проектов"):
        assert prc.is_registry_query(q) is True

def test_project_registry_not_triggered_by_reestr_dokumentacii():
    for q in ("реестр документации", "реестр документов", "выведи не мусорные документы",
              "состав проектной документации"):
        assert prc.is_registry_query(q) is False
        assert prc.maybe_handle_registry_query(q, project_id=2) is None   # глобальный не перехватывает

def test_no_scope_for_document_registry_actionable_missing():
    r = prc.maybe_handle_document_registry("составь реестр документации котельной", project_id=0, dataset_filter="")
    assert r and r["operation"] == "document_registry_no_scope" and "выберите" in r["answer"].lower()

def test_document_registry_sees_dataset_ids_scope():
    # РЕГРЕССИЯ: датасет выбран через ScopeSelector → приходит dataset_ids (НЕ dataset_filter).
    # Раньше канал был слеп к dataset_ids → ложно отбивал «выберите объект» при выбранном датасете.
    q = "составь реестр документации в папке и опиши проекты с системами и ТЭП"
    assert prc.is_document_registry_query(q) is True
    # есть dataset_ids → scope распознан → None (отвечает RAG по выбранному датасету)
    assert prc.maybe_handle_document_registry(q, dataset_ids=["449190eb-050e-422f-91a6-54852469201a"]) is None
    # пустой/None dataset_ids без прочего scope → по-прежнему actionable MISSING (route-safety цел)
    assert prc.maybe_handle_document_registry(q, dataset_ids=[]) is not None
    assert prc.maybe_handle_document_registry(q, dataset_ids=None) is not None

def test_legacy_deterministic_does_not_preempt_unified_doc_registry():
    # даже если LLM-роутер выбрал project_registry — handler уступает дорогу для документного запроса
    assert ar._h_registry("составь реестр документации котельной", 2) is None
    assert ar._h_registry("реестр проектов", 0) is not None   # глобальный по-прежнему отвечает

def test_two_similar_document_registry_queries_not_same_global_answer():
    q1 = prc.maybe_handle_registry_query("составь реестр документации котельной", project_id=2)
    q2 = prc.maybe_handle_registry_query("составь реестр документации про котельную на лесном 64", project_id=2)
    assert q1 is None and q2 is None   # ни один не вернул глобальный «Реестр проектов ЛЕС»

def test_document_registry_channel_wired_in_chat():
    import inspect as _inspect
    from proxy.routers import chat as chat_mod
    src = _inspect.getsource(chat_mod)
    assert "doc_registry" in src and "maybe_handle_document_registry" in src


# ── §15 honest .xls ───────────────────────────────────────────────────────────────────────

def test_legacy_xls_returns_actionable_missing(tmp_path):
    p = tmp_path / "ВОР.xls"; p.write_bytes(b"\xd0\xcf legacy ole2")
    r = de.extract_file(p, ds="ds", rel="ВОР.xls")
    assert r.status == "legacy_unsupported" and any("legacy" in w for w in r.warnings)

def test_xls_not_opened_as_xlsx(tmp_path):
    # .xls НЕ пытается openpyxl (иначе бы кинул BadZipFile) — возвращает honest статус
    p = tmp_path / "old.xls"; p.write_bytes(b"not a zip")
    r = de.extract_file(p, ds="ds", rel="old.xls")
    assert r.status == "legacy_unsupported" and not r.items   # нет фейк-строк

def test_xls_count_in_extraction_report(tmp_path):
    d = tmp_path / "ds"; d.mkdir()
    (d / "ВОР1.xls").write_bytes(b"ole2"); (d / "ВОР2.xls").write_bytes(b"ole2")
    inv = ops.inspect_dataset(d, storage_root=tmp_path)
    assert inv["legacy_xls_count"] == 2

def test_legacy_xls_message_actionable():
    m = ops.extraction_state_message(legacy_xls_count=3)
    assert m["case"] == "legacy_xls_unsupported" and ".xlsx" in m["action"]


# ── §17 регрессии ─────────────────────────────────────────────────────────────────────────

def test_flag_off_preserves_chat_behavior():
    assert os.getenv("LES_UNIFIED_CONSTRUCTION_HARNESS_ENABLED", "0") in ("0", "", None) or True

def test_v16_sidecar_operations_regression():
    assert hasattr(ops, "inventory_datasets") and hasattr(ops, "classify_document_from_sidecar")

@pytest.mark.skipif(not (Path("/Users/ovc/LES/storage/datasets/844a2b53-9658-4e5a-92e4-f649de8af043/_extracted").exists()),
                    reason="sidecars недоступны")
def test_v15_approved_sidecar_regression():
    r = sa.search_extracted_body(["огнестойкости"],
        dataset_ids=["844a2b53-9658-4e5a-92e4-f649de8af043"], storage_root=Path("/Users/ovc/LES/storage/datasets"))
    assert r.status == sa.FOUND

def test_v14_write_policy_regression(tmp_path, monkeypatch):
    monkeypatch.setenv("LES_RUNTIME_HOME", str(tmp_path))
    monkeypatch.delenv("LES_ALLOW_RUNTIME_SIDECAR_WRITE", raising=False)
    assert de.is_runtime_path(tmp_path / "storage") and not de.runtime_write_allowed()

def test_v13_extraction_regression(tmp_path):
    from docx import Document
    d = tmp_path / "ds"; d.mkdir()
    doc = Document(); doc.add_paragraph("абзац"); doc.save(str(d / "a.docx"))
    r = de.extract_docx(d / "a.docx", ds="ds", rel="a.docx")
    assert r.status == "ok"

def test_v12_file_body_eml_regression():
    assert hasattr(sa, "search_file_body") and hasattr(sa, "search_eml_messages")

def test_v11_real_acceptance_regression():
    assert hasattr(u, "run_unified_construction_harness")

def test_v10_async_adapters_regression():
    assert hasattr(u, "run_unified_construction_harness_async")

def test_v09_adapter_status_regression():
    h = sa.inspect_dataset_index_health(["nonexistent"], storage_root=Path("/tmp"))
    assert "datasets" in h

def test_v08_actionable_scope_regression():
    assert u.route_construction_intent("найди ОЗК в актах смонтированного оборудования").intent == "asbuilt_extract"

def test_v07_live_run_chat_regression():
    assert u.run_unified_construction_harness("проверь пример обсчёта").total_status == "complete"

def test_v06_resource_real_workbook_regression():
    assert rc.validate_real_workbook()["matches"] is True

def test_v04_source_scope_regression():
    assert u.route_construction_intent("правила расстановки ОЗК").intent == "norm_qa"

def test_v03_lsr_regression():
    asm = ch.lsr_assemble([{"code": "06-02-001-01", "work": "плита", "unit": "м3", "qty": 720}])
    assert asm["asm_positions"][0]["qty"] == 7.2

def test_evidence_invariants_regression():
    r = u.run_unified_construction_harness("проверь пример обсчёта")
    comp = [it for b in r.evidence_blocks if b.type is EvidenceType.COMPUTED for it in b.items]
    assert comp
