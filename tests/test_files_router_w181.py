"""W18.1 — отдача файлов/структуры: path-guard, дерево, текст."""
import pytest

from fastapi import HTTPException

from proxy.routers import files as files_router


@pytest.fixture()
def root(tmp_path, monkeypatch):
    import fitz
    import openpyxl
    from docx import Document

    (tmp_path / "NTD").mkdir()
    (tmp_path / "NTD" / "note.md").write_text("# Привет\nтекст", encoding="utf-8")
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "LES PDF evidence")
    page.draw_rect(fitz.Rect(68, 56, 230, 84), color=(0, 0.5, 0))
    pdf.save(tmp_path / "NTD" / "doc.pdf")
    pdf.close()
    word = Document()
    word.add_heading("Проект", 1)
    word.add_paragraph("Проверяемый абзац")
    word.save(tmp_path / "NTD" / "report.docx")
    workbook = openpyxl.Workbook()
    workbook.active.title = "ВОР"
    workbook.active.append(["Работа", "Количество"])
    workbook.active.append(["Монтаж", 12.5])
    workbook.create_sheet("Итоги").append(["Всего", "=SUM(ВОР!B2:B2)"])
    workbook.save(tmp_path / "NTD" / "register.xlsx")
    workbook.close()
    (tmp_path / ".hidden").write_text("secret")
    monkeypatch.setattr(files_router, "_ROOT", tmp_path)
    return tmp_path


def test_safe_blocks_traversal(root):
    with pytest.raises(HTTPException) as e:
        files_router._safe("../../etc/passwd")
    assert e.value.status_code == 400


def test_safe_allows_inside(root):
    p = files_router._safe("NTD/note.md")
    assert p.name == "note.md"


def test_safe_accepts_rag_content_prefixed_source_ref(root):
    p = files_router._safe("RAG_Content/NTD/note.md")
    assert p == root / "NTD" / "note.md"


@pytest.mark.asyncio
async def test_tree_lists_dirs_first_skips_hidden(root):
    tree = await files_router.rag_tree(path="", depth=2, _user=object())
    assert tree["dir"] is True
    names = [c["name"] for c in tree["children"]]
    assert "NTD" in names
    assert ".hidden" not in names  # скрытые не показываем
    ntd = next(c for c in tree["children"] if c["name"] == "NTD")
    files = {c["name"]: c for c in ntd["children"]}
    assert files["note.md"]["dir"] is False
    assert files["note.md"]["path"] == "NTD/note.md"


@pytest.mark.asyncio
async def test_file_text_returns_content(root):
    res = await files_router.rag_file_text(path="NTD/note.md", _user=object())
    assert res["language"] == "md"
    assert "Привет" in res["content"]


@pytest.mark.asyncio
async def test_file_text_rejects_binary(root):
    with pytest.raises(HTTPException) as e:
        await files_router.rag_file_text(path="NTD/doc.pdf", _user=object())
    assert e.value.status_code == 415  # бинарь → через /file/raw


@pytest.mark.asyncio
async def test_file_text_404_missing(root):
    with pytest.raises(HTTPException) as e:
        await files_router.rag_file_text(path="NTD/nope.md", _user=object())
    assert e.value.status_code == 404


@pytest.mark.asyncio
async def test_file_raw_serves_pdf(root):
    resp = await files_router.rag_file_raw(path="NTD/doc.pdf", _user=object())
    assert str(resp.path).endswith("doc.pdf")


@pytest.mark.asyncio
async def test_pdf_info_preview_and_unified_viewer(root):
    info = await files_router.rag_file_pdf_info(path="NTD/doc.pdf", _user=object())
    assert info["page_count"] == 1
    preview = await files_router.rag_file_pdf_preview(
        path="NTD/doc.pdf", page=1, width=640, highlight_bbox="68,56,230,84", _user=object()
    )
    assert preview.media_type == "image/png"
    assert preview.body.startswith(b"\x89PNG")
    viewer = await files_router.rag_file_viewer(
        path="NTD/doc.pdf", page=1, bbox="68,56,230,84", locator="p1", sheet="", _user=object()
    )
    assert b"list.pdf_viewer.v1" in viewer.body
    assert b"pdf-preview" in viewer.body
    assert str(root).encode() not in viewer.body


@pytest.mark.asyncio
async def test_unified_viewer_opens_docx_and_xlsx(root):
    word = await files_router.rag_file_viewer(
        path="NTD/report.docx", page=1, bbox="", locator="para1", sheet="", _user=object()
    )
    assert b"list.file_viewer.v1" in word.body
    assert "Проверяемый абзац".encode() in word.body
    assert b"evidence-hit" in word.body

    excel = await files_router.rag_file_viewer(
        path="NTD/register.xlsx", page=1, bbox="", locator="ВОР!R2", sheet="ВОР", _user=object()
    )
    assert "Монтаж".encode() in excel.body
    assert 'data-sheet="ВОР"'.encode() in excel.body
    assert b"evidence-hit" in excel.body


@pytest.mark.asyncio
async def test_unified_viewer_rejects_unknown_binary(root):
    (root / "NTD" / "model.bin").write_bytes(b"binary")
    with pytest.raises(HTTPException) as exc:
        await files_router.rag_file_viewer(
            path="NTD/model.bin", page=1, bbox="", locator="", sheet="", _user=object()
        )
    assert exc.value.status_code == 415


# ── Мульти-корень: внешние корни индексации по ссылке (ADR-12) ──

@pytest.fixture()
def ext_root(root, tmp_path, monkeypatch):
    ext = tmp_path / "ext_src"
    (ext / "Котельная" / "РД").mkdir(parents=True)
    (ext / "Котельная" / "note.md").write_text("# схема", encoding="utf-8")
    monkeypatch.setenv("LES_EXTERNAL_SOURCE_ROOTS", str(ext))
    return ext


@pytest.mark.asyncio
async def test_tree_top_lists_internal_and_external_roots(ext_root):
    tree = await files_router.rag_tree(path="", depth=2, _user=object())
    # синтетический супер-корень с детьми-корнями
    names = {c["name"] for c in tree["children"]}
    assert "RAG_Content" in names
    assert "ext_src" in names


@pytest.mark.asyncio
async def test_external_path_prefixed_and_resolves(ext_root):
    tree = await files_router.rag_tree(path="ext_src::", depth=2, _user=object())
    kot = next(c for c in tree["children"] if c["name"] == "Котельная")
    assert kot["path"].startswith("ext_src::")
    res = await files_router.rag_file_text(path="ext_src::Котельная/note.md", _user=object())
    assert "схема" in res["content"]


def test_external_traversal_blocked(ext_root):
    with pytest.raises(HTTPException) as e:
        files_router._safe("ext_src::../../etc/passwd")
    assert e.value.status_code == 400


def test_unknown_root_rejected(ext_root):
    with pytest.raises(HTTPException) as e:
        files_router._safe("nope::secret")
    assert e.value.status_code == 400
